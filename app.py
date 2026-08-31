import streamlit as st
import os
import datetime
import io
import asyncio
import time
import base64
import pytz
import sqlite3
import urllib.parse
import hashlib
import secrets 
import pandas as pd
from dataclasses import dataclass
from contextlib import contextmanager
from streamlit_mic_recorder import mic_recorder
from duckduckgo_search import DDGS
from PIL import Image
from pypdf import PdfReader
import docx
from docx import Document
from groq import Groq
import json
import streamlit.components.v1 as components
import requests
import concurrent.futures
from googlesearch import search as google_search
import requests
from bs4 import BeautifulSoup
from RestrictedPython import compile_restricted, safe_builtins
from RestrictedPython.PrintCollector import PrintCollector
import re
import json
import concurrent.futures
import numpy as np
import json
import math
import datetime
import concurrent.futures
import httpx
from bs4 import BeautifulSoup
import datetime
import json
import math
import urllib.parse
import concurrent.futures
from sentence_transformers import SentenceTransformer
import aiohttp

ALLOWED_MODELS = [
    "openai/gpt-oss-20b",
    "qwen/qwen3.8-27b",
    "groq/compound",
    "openai/gpt-oss-120b"
]

selected_model = st.sidebar.selectbox(
    "Válassz modellt:",
    options=ALLOWED_MODELS,
    index=0
)

@st.cache_data(ttl=3600)
def cached_tool_query(query: str, tool_type: str):
    pass


def magyar_szoto_normalizalo(text: str) -> list[str]:
    """
    Kiszűri a magyar ragokat és toldalékokat a pontosabb kulcsszó-egyeztetéshez.
    """
    words = re.findall(r'\b\w+\b', text.lower())
    clean_tokens = []
    # Gyakoribb magyar ragok / toldalékok levágása (heurisztikus stemmer)
    suffixes = [
        'ban', 'ben', 'nak', 'nek', 'val', 'vel', 'ból', 'ből', 'ról', 'ről',
        'hoz', 'hez', 'höz', 'ig', 'ért', 'ba', 'be', 'ra', 're', 'at', 'et',
        'ot', 'öt', 'k', 'ak', 'ek', 'ok', 'ök', 'ja', 'je', 'ai', 'ei'
    ]
    for w in words:
        if w in HUNGARIAN_STOPWORDS or len(w) <= 2:
            continue
        stemmed = w
        for suf in suffixes:
            if stemmed.endswith(suf) and len(stemmed) - len(suf) >= 3:
                stemmed = stemmed[:-len(suf)]
                break
        clean_tokens.append(stemmed)
    return clean_tokens

def hibrid_rrf_rangsorolas(vector_results: list, bm25_results: list, k: int = 60) -> list:
    """
    Reciprocal Rank Fusion (RRF) egyesíti a szemantikus vektoros és a BM25 kulcsszavas találatokat.
    """
    scores = {}
    
    for rank, doc in enumerate(vector_results):
        doc_id = doc['url']
        scores[doc_id] = scores.get(doc_id, 0) + 1.0 / (k + rank + 1)
        
    for rank, doc in enumerate(bm25_results):
        doc_id = doc['url']
        scores[doc_id] = scores.get(doc_id, 0) + 1.0 / (k + rank + 1)
        
    all_docs = {d['url']: d for d in vector_results + bm25_results}
    sorted_docs = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    return [all_docs[doc_id] for doc_id, _ in sorted_docs]

@st.cache_resource
def get_embedding_model():
    # Ingyenes, gyors, és érti a magyar nyelvet
    return SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')

try:
    from googlesearch import search as google_search
    HAS_GOOGLE_SEARCH = True
except ImportError:
    HAS_GOOGLE_SEARCH = False

try:
    import trafilatura
    HAS_TRAFILATURA = True
except ImportError:
    HAS_TRAFILATURA = False

HUNGARIAN_STOPWORDS = {
    "a", "az", "egy", "be", "ki", "le", "fel", "meg", "el", "at", "es", "hogy", 
    "nem", "sem", "vagy", "is", "csak", "mint", "volt", "lesz", "cikk", "alatt",
    "van", "vannak", "ma", "majd", "mert", "ha", "de", "rabs", "mely", "amely",
    "ebben", "ebbol", "arról", "melyek", "szerint", "után", "során"
}

def optimalizal_keresesi_kifejezeseket(client, felhasznalo_kerdese: str, model_name: str = None) -> list[str]:
    most = datetime.datetime.now()
    aktualis_datum = most.strftime("%Y-%m-%d")
    aktualis_ev = most.year
    
    osszes_modell = fetch_groq_models(GROQ_API_KEY)
    szurt_modellek = [m for m in osszes_modell if m in ALLOWED_MODELS] if osszes_modell else ALLOWED_MODELS
    
    if not model_name or model_name not in szurt_modellek:
        model_name = szurt_modellek[0]

    try:
        prompt = f"""
        Ma {aktualis_datum} van ({aktualis_ev}. év), de ezt ne emlegetsd csak vedd figyelembe a válaszadásnál.
        Hozz létre pontosan 3 eltérő, rövid és időszerű keresőkifejezést webes kereséshez a következő kérdésből.
        Ha a kérdés friss eseményre utal, építsd be a(z) {aktualis_ev} évet!

        Kizárólag egy érvényes JSON tömböt adj vissza stringekkel!
        Példa: ["kifejezés 1", "kifejezés 2", "kifejezés 3"]

        Kérdés: {felhasznalo_kerdese}
        """

        def safe_completion(client, messages, chosen_model, allowed_list):
            models_to_try = [chosen_model] + [m for m in allowed_list if m != chosen_model]
            for model in models_to_try:
                try:
                    return client.chat.completions.create(
                        model=model,
                        messages=messages
                    )
                except Exception:
                    continue
            raise RuntimeError("Egyik engedélyezett modell sem válaszolt az API-n keresztül.")

        response = safe_completion(client, [{"role": "user", "content": prompt}], model_name, szurt_modellek)
        content = response.choices[0].message.content.strip()
        match = re.search(r'\[.*\]', content, re.DOTALL)
        if match:
            queries = json.loads(match.group(0))
            if isinstance(queries, list) and len(queries) > 0:
                return queries[:3]
    except Exception:
        pass
    return [felhasznalo_kerdese]

async def kereses_es_szures(kifejezesek: list[str], max_talalat_per_kifejezes=2) -> list[str]:
    osszes_url = set()
    async with DDGS() as ddgs:
        feladatok = [ddgs.atext(k, max_results=max_talalat_per_kifejezes) for k in kifejezesek]
        eredmenyek = await asyncio.gather(*feladatok, return_exceptions=True)
        
        for eredmeny_lista in eredmenyek:
            if isinstance(eredmeny_lista, list):
                for talalat in eredmeny_lista:
                    if "href" in talalat:
                        osszes_url.add(talalat["href"])
    return list(osszes_url)

async def tiszta_szoveg_kinyerese(url: str, session: aiohttp.ClientSession) -> str:
    fejlec = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }
    try:
        async with session.get(url, headers=fejlec, timeout=6) as response:
            if response.status != 200:
                return ""
            html = await response.text()
            
            soup = BeautifulSoup(html, "html.parser")
            for zaj in soup(["script", "style", "header", "footer", "nav", "aside"]):
                zaj.extract()
                
            szoveg = soup.get_text(separator=" ", strip=True)
            
            tiszta_szoveg = " ".join(szoveg.split())
            if not tiszta_szoveg:
                return ""
                
            return f"FORRÁS ({url}):\n{tiszta_szoveg[:3500]}...\n"
    except Exception:
        return ""

async def webes_kontextus_generalasa(kifejezesek: list[str]) -> str:
    url_lista = await kereses_es_szures(kifejezesek)
    if not url_lista:
        return "Nem sikerült releváns webes találatokat lekérni."
    
    async with aiohttp.ClientSession() as session:
        scraping_feladatok = [tiszta_szoveg_kinyerese(url, session) for url in url_lista]
        szovegek = await asyncio.gather(*scraping_feladatok)
        
    vegso_kontextus = "\n---\n".join(filter(bool, szovegek))
    return vegso_kontextus


def bing_ingyenes_kereses(query: str, max_results: int = 5) -> list[dict]:
    """
    Ingyenes Bing Webes Keresőmotor (Scraper API-kulcs nélkül).
    Valódi friss webes találatokat gyűjt be Wikipédia használata nélkül.
    """
    results = []
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36',
        'Accept-Language': 'hu-HU,hu;q=0.9,en-US;q=0.8,en;q=0.7'
    }
    encoded_query = urllib.parse.quote(query)
    search_url = f"https://www.bing.com/search?q={encoded_query}"
    
    try:
        with httpx.Client(timeout=5, follow_redirects=True, headers=headers) as client:
            resp = client.get(search_url)
            if resp.status_code == 200:
                soup = BeautifulSoup(resp.text, 'html.parser')
                items = soup.find_all('li', class_='b_algo')
                for item in items[:max_results]:
                    h2 = item.find('h2')
                    if not h2:
                        continue
                    a_tag = h2.find('a')
                    if not a_tag or not a_tag.get('href'):
                        continue
                    
                    title = a_tag.get_text(strip=True)
                    link = a_tag['href']
                    snippet_elem = item.find('p') or item.find('div', class_='b_caption')
                    snippet = snippet_elem.get_text(strip=True) if snippet_elem else ""
                    
                    if link.startswith('http'):
                        results.append({
                            'title': title,
                            'url': link,
                            'snippet': snippet
                        })
    except Exception:
        pass
    return results


def letolt_es_tisztit_html(url: str, timeout: int = 6) -> str:
    """
    Biztonságos és zajmentes weboldal-letöltés trafilatura / BeautifulSoup fallbackkel.
    """
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
        'Accept-Language': 'hu-HU,hu;q=0.9,en-US;q=0.8,en;q=0.7'
    }
    try:
        with httpx.Client(timeout=timeout, follow_redirects=True, headers=headers, verify=False) as http_client:
            resp = http_client.get(url)
            
            content_type = resp.headers.get("content-type", "").lower()
            if "text/html" not in content_type and "text/plain" not in content_type:
                return ""

            if resp.status_code != 200:
                return ""
            
            html_content = resp.text

            if HAS_TRAFILATURA:
                extracted = trafilatura.extract(
                    html_content, 
                    include_comments=False, 
                    include_tables=True,
                    favor_precision=True
                )
                if extracted and len(extracted.strip()) > 120:
                    return extracted.strip()

            soup = BeautifulSoup(html_content, 'html.parser')
            for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form", "iframe", "noscript", "svg"]):
                tag.extract()
            
            main_content = soup.find('main') or soup.find('article') or soup.find('body')
            if main_content:
                paragraphs = [elem.get_text(" ", strip=True) for elem in main_content.find_all(['p', 'h1', 'h2', 'h3', 'li', 'td'])]
                full_text = "\n".join([p for p in paragraphs if len(p) > 25])
                return full_text
    except Exception:
        pass
    return ""


def szamits_bm25_n_gram_pont(query: str, title: str, text: str) -> float:
    """
    Továbbfejlesztett BM25 + Bigram / Phrase Matching algoritmus.
    """
    def tokenize(s: str) -> list[str]:
        words = re.findall(r'\w+', s.lower())
        return [w for w in words if w not in HUNGARIAN_STOPWORDS and len(w) > 2]

    q_tokens = tokenize(query)
    if not q_tokens:
        return 0.0

    t_tokens = tokenize(text)
    title_tokens = tokenize(title)

    if not t_tokens:
        return 0.0

    k1 = 1.2
    b = 0.75
    avg_doc_len = 250
    doc_len = len(t_tokens)

    score = 0.0
    for token in set(q_tokens):
        tf = t_tokens.count(token)
        if tf == 0:
            continue
        tf_score = (tf * (k1 + 1)) / (tf + k1 * (1 - b + b * (doc_len / avg_doc_len)))
        score += tf_score

        if token in title_tokens:
            score += 3.0

    if len(q_tokens) >= 2:
        for i in range(len(q_tokens) - 1):
            bigram = f"{q_tokens[i]} {q_tokens[i+1]}"
            if bigram in text.lower():
                score += 4.5

    return score


def kiemel_szemantikus_ablakokat_hibrid(query: str, full_text: str, max_chars: int = 2000) -> str:
    """
    Továbbfejlesztett csúsztatott ablakos algoritmus, amely a BM25 kulcsszavas
    és a MiniLM szemantikus vektoros egyezéseket (Hibrid) ötvözi.
    """
    sentences = re.split(r'(?<=[.!?])\s+', full_text.strip())
    if not sentences:
        return full_text[:max_chars]

    windows = []
    window_size = 3
    for i in range(0, max(1, len(sentences) - window_size + 1)):
        chunk = " ".join(sentences[i:i + window_size])
        if len(chunk) > 40:
            windows.append(chunk)

    if not windows:
        return full_text[:max_chars]

    embedder = get_embedding_model()
    query_vector = embedder.encode(query)
    window_vectors = embedder.encode(windows)

    scored_windows = []
    for idx, chunk in enumerate(windows):
        bm25_score = szamits_bm25_n_gram_pont(query, "", chunk)
        
        chunk_vector = window_vectors[idx]
        cos_sim = np.dot(query_vector, chunk_vector) / (np.linalg.norm(query_vector) * np.linalg.norm(chunk_vector))
        semantic_score = float(cos_sim) if not np.isnan(cos_sim) else 0.0

        final_score = (semantic_score * 15.0) + bm25_score
        
        scored_windows.append((final_score, chunk))

    scored_windows.sort(key=lambda x: x[0], reverse=True)

    selected_chunks = []
    current_len = 0
    seen_chunks = set()

    for score, chunk in scored_windows:
        if chunk in seen_chunks:
            continue
        if score < 2.5: 
            continue
            
        if current_len + len(chunk) > max_chars:
            break
            
        selected_chunks.append(chunk)
        seen_chunks.add(chunk)
        current_len += len(chunk)

    return "\n\n[...] ".join(selected_chunks) if selected_chunks else full_text[:max_chars]


def hajzsalpontos_web_kereses(client, query: str, max_sources: int = 5) -> str:
    """
    3-szoros Ingyenes Hibrid Kereső Motor (DuckDuckGo + Google Search + Bing Scraper)
    Wikipedia NÉLKÜL, 100%-ban friss és ingyenes webes találatokkal.
    """
    search_queries = optimalizal_keresesi_kifejezeseket(client, query)
    
    raw_results = []
    seen_urls = set()

    for q_idx, sq in enumerate(search_queries):
        # 1. MOTOR: DuckDuckGo Keresés
        try:
            with DDGS() as ddgs:
                ddg_res = list(ddgs.text(sq, max_results=4))
                for rank, r in enumerate(ddg_res):
                    url = r.get('href')
                    if url and url not in seen_urls:
                        seen_urls.add(url)
                        raw_results.append({
                            'title': r.get('title', ''),
                            'url': url,
                            'snippet': r.get('body', ''),
                            'initial_rank': rank,
                            'query_idx': q_idx
                        })
        except Exception:
            pass

        if HAS_GOOGLE_SEARCH and len(raw_results) < 6:
            try:
                g_urls = list(google_search(sq, num_results=3, sleep_interval=0.5))
                for rank, g_url in enumerate(g_urls):
                    if g_url and g_url not in seen_urls:
                        seen_urls.add(g_url)
                        raw_results.append({
                            'title': g_url,
                            'url': g_url,
                            'snippet': '',
                            'initial_rank': rank + 2,
                            'query_idx': q_idx
                        })
            except Exception:
                pass

        if len(raw_results) < 8:
            bing_res = bing_ingyenes_kereses(sq, max_results=4)
            for rank, b_item in enumerate(bing_res):
                url = b_item['url']
                if url and url not in seen_urls:
                    seen_urls.add(url)
                    raw_results.append({
                        'title': b_item['title'],
                        'url': url,
                        'snippet': b_item['snippet'],
                        'initial_rank': rank + 1,
                        'query_idx': q_idx
                    })

    if not raw_results:
        return "Nem találtam releváns friss információt a weben a megadott kérdésre."

    fetched_data = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=8) as executor:
        future_to_item = {
            executor.submit(letolt_es_tisztit_html, item['url']): item 
            for item in raw_results[:10]
        }
        for future in concurrent.futures.as_completed(future_to_item):
            item = future_to_item[future]
            try:
                page_text = future.result()
            except Exception:
                page_text = ""

            if len(page_text) > 120:
                content = kiemel_szemantikus_ablakokat_hibrid(query, page_text, max_chars=1800)
            else:
                content = item['snippet']

            bm25_score = szamits_bm25_n_gram_pont(query, item['title'], content)
            rrf_score = 1.0 / (60 + item['initial_rank'])
            
            final_score = bm25_score + (rrf_score * 15.0)

            fetched_data.append({
                'title': item['title'],
                'url': item['url'],
                'content': content,
                'score': final_score
            })

    fetched_data.sort(key=lambda x: x['score'], reverse=True)
    top_sources = fetched_data[:max_sources]

    kontextus_blokkok = []
    for idx, src in enumerate(top_sources, 1):
        blokk = (
            f"FORRÁS [{idx}]:\n"
            f"Cím: {src['title']}\n"
            f"URL: {src['url']}\n"
            f"Tartalom:\n{src['content']}\n"
            f"----------------------------------------"
        )
        kontextus_blokkok.append(blokk)

    return "\n\n".join(kontextus_blokkok)


def generald_a_hajszalpontos_valaszt(client, felhasznalo_kerdese: str, web_kontextus: str = "", doc_kontextus: str = "", model_name: str = None):
    most = datetime.datetime.now()
    aktualis_datum = most.strftime("%Y. %B %d.")

    if not model_name:
        available = fetch_groq_models(GROQ_API_KEY)
        model_name = available[0] if available else "qwen/qwen3.8-27b"

    system_prompt = f"""
Te egy prémium szintű, tényalapú intelligens asszisztens vagy.
A mai dátum: {aktualis_datum}.

utasítások a PONTOSÁG ÉS MEGBÍZHATÓSÁG ÉRDEKÉBEN:
1. **Gondolkodási folyamat (Chain-of-Thought):** Mielőtt megadnád a végső választ, hajtsd végre a következő belső lépéseket:
   - Elemezd a kérdés pontos célját és a rendelkezésre álló kontextust!
   - Különítsd el az igazolt tényeket az esetleges ellentmondásoktól!
   - Ha matematikai, kódolási vagy logikai feladatról van szó, lépésről lépésre számolj/gondolkodj!

2. **Források és Hivatkozások:**
   - Amennyiben webes keresési vagy dokumentum kontextus áll rendelkezésre, szigorúan használd a kattintható Markdown hivatkozásokat! Példa: `[1](https://forras.com)`.
   - Ha a kapott kontextus hiányos, de a kérdés általános műveltségi/logikai/kódolási jellegű, használd a saját, mély logikai tudásodat, de jelezd a bizonytalansági tényezőket!

3. **Stílus:**
   - Legyél lényegre törő, áttekinthető, strukturált és 100%-ig precíz.
"""

    response = client.chat.completions.create(
        model=model_name,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": felhasznalo_kerdese}
        ],
        temperature=0.1,
        max_tokens=3000
    )

    return response.choices[0].message.content




def render_gps_navigation(dest_name="", dest_lat=None, dest_lng=None):
    """
    Dinamikus GPS térkép beágyazása:
    - Lekéri a böngészőből a felhasználó aktuális GPS koordinátáit.
    - Ráfókuszál a felhasználóra (kék pulzáló pont).
    - Ha meg van adva célállomás (dest_lat, dest_lng), kirajzolja az útvonalat.
    """
    
    dest_data_json = json.dumps({
        "name": dest_name,
        "lat": dest_lat,
        "lng": dest_lng
    })

    html_code = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8" />
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css" />
        <script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
        
        <link rel="stylesheet" href="https://unpkg.com/leaflet-routing-machine@latest/dist/leaflet-routing-machine.css" />
        <script src="https://unpkg.com/leaflet-routing-machine@latest/dist/leaflet-routing-machine.js"></script>

        <style>
            body { margin: 0; padding: 0; font-family: Arial, sans-serif; }
            #map { height: 480px; width: 100%; border-radius: 12px; box-shadow: 0 4px 12px rgba(0,0,0,0.15); }
            
            .gps-status {
                padding: 10px 14px;
                background-color: #f0f2f6;
                border-left: 4px solid #007bff;
                border-radius: 6px;
                margin-bottom: 10px;
                font-size: 14px;
                font-weight: bold;
                color: #333;
            }

            /* Pulzáló kék GPS jelölő a felhasználó pozíciójához */
            .user-gps-dot {
                width: 18px;
                height: 18px;
                background-color: #007bff;
                border: 3px solid #ffffff;
                border-radius: 50%;
                box-shadow: 0 0 10px rgba(0, 123, 255, 0.9);
                animation: pulse 1.6s infinite;
            }

            @keyframes pulse {
                0% { box-shadow: 0 0 0 0 rgba(0, 123, 255, 0.7); }
                70% { box-shadow: 0 0 0 14px rgba(0, 123, 255, 0); }
                100% { box-shadow: 0 0 0 0 rgba(0, 123, 255, 0); }
            }
        </style>
    </head>
    <body>
        <div id="status" class="gps-status"> GPS kapcsolat keresése...</div>
        <div id="map"></div>

        <script>
            const destData = __DEST_DATA_JSON__;
            const statusDiv = document.getElementById('status');

            // Alapértelmezett térkép (Budapest központ fallback)
            const map = L.map('map').setView([47.4979, 19.0402], 13);

            L.tileLayer('https://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png', {
                maxZoom: 19,
                attribution: '© OpenStreetMap'
            }).addTo(map);

            //  Böngésző GPS Helymeghatározása
            if ("geolocation" in navigator) {
                navigator.geolocation.getCurrentPosition(
                    (position) => {
                        const userLat = position.coords.latitude;
                        const userLng = position.coords.longitude;

                        statusDiv.innerHTML = "✅ GPS pozíció beérkezve! Ráállás a helyzetedre...";

                        // Kék GPS ikon létrehozása
                        const userIcon = L.divIcon({
                            className: 'user-gps-dot',
                            iconSize: [18, 18],
                            iconAnchor: [9, 9]
                        });

                        // Felhasználó megjelölése
                        L.marker([userLat, userLng], { icon: userIcon })
                         .addTo(map)
                         .bindPopup("<b> Az Ön jelenlegi pozíciója</b>")
                         .openPopup();

                        // GPS Fókusz a felhasználóra
                        map.setView([userLat, userLng], 15);

                        // 🏁 Ha van célállomás, útvonal kirajzolása
                        if (destData.lat && destData.lng) {
                            L.Routing.control({
                                waypoints: [
                                    L.latLng(userLat, userLng),
                                    L.latLng(destData.lat, destData.lng)
                                ],
                                router: L.Routing.osrmv1({
                                    serviceUrl: 'https://router.project-osrm.org/route/v1'
                                }),
                                routeWhileDragging: false,
                                show: true,
                                collapsible: true,
                                createMarker: function(i, wp, n) {
                                    if (i === 0) return null;
                                    return L.marker(wp.latLng).bindPopup("<b> Célállomás: " + (destData.name || "Cél") + "</b>");
                                }
                            }).addTo(map);

                            statusDiv.innerHTML = " Útvonal megtervezve a célállomáshoz: <b>" + (destData.name || "Cél") + "</b>";
                        }
                    },
                    (error) => {
                        console.error("GPS Hiba:", error);
                        statusDiv.innerHTML = "⚠️ Nem sikerült lekérni a GPS pozíciót. Kérjük engedélyezd a helymeghatározást a böngészőben!";
                    },
                    {
                        enableHighAccuracy: true,
                        timeout: 10000,
                        maximumAge: 0
                    }
                );
            } else {
                statusDiv.innerHTML = " A böngésződ nem támogatja a GPS helymeghatározást.";
            }
        </script>
    </body>
    </html>
    """
    
    html_code = html_code.replace("__DEST_DATA_JSON__", dest_data_json)
    components.html(html_code, height=530)

@dataclass(frozen=True)
class AppConfig:
    DB_FILE: str = "zoli_gpt_local.db"
    ADMIN_USERNAME: str = st.secrets.get("ADMIN_USERNAME", "default_admin_fallback") 
    TIMEZONE: str = "Europe/Budapest"
    PIXABAY_API_KEY: str = st.secrets.get("PIXABAY_API_KEY", "") 
    MAX_HISTORY_CHARS: int = 4000
    RAG_SIMI_THRESHOLD: float = 0.15
    CHUNK_SIZE: int = 800
    CHUNK_OVERLAP: int = 300
    HUNGARIAN_STOPWORDS = frozenset([
        "a", "az", "egy", "be", "ki", "le", "fel", "meg", "el", "at", "es", "hogy", 
        "nem", "sem", "vagy", "is", "csak", "mint", "volt", "lesz", "cikk", "alatt"
    ])

st.set_page_config(page_title="Zoli GPT ", page_icon="🚭", layout="centered")

cfg = AppConfig()

if "generating" not in st.session_state:
    st.session_state.generating = False
else:
    st.session_state.generating = False

if "logged_in_user" not in st.session_state:
    st.session_state.logged_in_user = None
if "login_attempts" not in st.session_state:
    st.session_state.login_attempts = 0
if "lockout_until" not in st.session_state:
    st.session_state.lockout_until = 0

def hash_password(password: str, salt: str = None) -> tuple:
    if salt is None:
        salt = secrets.token_hex(16)
    # 100 000 iterációs PBKDF2 védelem szivárványtáblák és brute-force ellen
    key = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt.encode('utf-8'), 100000)
    return key.hex(), salt

def analyze_image_with_qwen(client, image_bytes: bytes, prompt: str = "Elemzed a képet és írd le részletesen, mi látható rajta!", model_name: str = "qwen/qwen3.8-27b") -> str:
    try:
        base64_image = base64.b64encode(image_bytes).decode('utf-8')
        messages = [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                ]
            }
        ]
        response = client.chat.completions.create(model=model_name, messages=messages, max_tokens=1000)
        return response.choices[0].message.content
    except Exception as e:
        return f"Hiba a kép elemzése során: {e}"

class DatabaseRepository:
    def __init__(self, db_file: str):
        self.db_file = db_file
        self._init_schema()

    @contextmanager
    def _get_connection(self):
        conn = sqlite3.connect(self.db_file, check_same_thread=False, timeout=10.0)
        try: yield conn
        finally: conn.close()

    def _init_schema(self):
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''CREATE TABLE IF NOT EXISTS users (username TEXT PRIMARY KEY, password_hash TEXT)''')
            cursor.execute('''CREATE TABLE IF NOT EXISTS chat_history (id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT, role TEXT, content TEXT, type TEXT, caption TEXT, timestamp TEXT, thread_id TEXT DEFAULT 'default')''')
            cursor.execute('''CREATE TABLE IF NOT EXISTS document_vectors (id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT, doc_name TEXT, chunk_text TEXT, embedding BLOB, file_size TEXT)''')
            cursor.execute('''CREATE TABLE IF NOT EXISTS latency_logs (id INTEGER PRIMARY KEY AUTOINCREMENT, duration REAL, timestamp TEXT)''')
            cursor.execute('''CREATE TABLE IF NOT EXISTS token_logs (id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT, tokens INTEGER, cost REAL, timestamp TEXT)''')
            cursor.execute('''CREATE TABLE IF NOT EXISTS system_alerts (id INTEGER PRIMARY KEY AUTOINCREMENT, message TEXT, timestamp TEXT)''')
            
            try: cursor.execute("ALTER TABLE chat_history ADD COLUMN thread_id TEXT DEFAULT 'default'")
            except sqlite3.OperationalError: pass
            try: cursor.execute("ALTER TABLE token_logs ADD COLUMN username TEXT")
            except sqlite3.OperationalError: pass
            try: cursor.execute("ALTER TABLE token_logs ADD COLUMN tokens INTEGER")
            except sqlite3.OperationalError: pass
            try: cursor.execute("ALTER TABLE token_logs ADD COLUMN cost REAL")
            except sqlite3.OperationalError: pass
            try: cursor.execute("ALTER TABLE token_logs ADD COLUMN timestamp TEXT")
            except sqlite3.OperationalError: pass
            
            try: cursor.execute("ALTER TABLE users ADD COLUMN salt TEXT")
            except sqlite3.OperationalError: pass
            
            conn.commit()

    def register_user(self, username: str, password_raw: str) -> bool:
        pwd_hash, salt = hash_password(password_raw)
        with self._get_connection() as conn:
            cursor = conn.cursor()
            try:
                cursor.execute("INSERT INTO users (username, password_hash, salt) VALUES (?, ?, ?)", (username, pwd_hash, salt))
                conn.commit()
                return True
            except sqlite3.IntegrityError:
                return False

    def verify_user(self, username: str, password_raw: str) -> bool:
        with self._get_connection() as conn:
            cursor = conn.cursor()
            try:
                cursor.execute("SELECT password_hash, salt FROM users WHERE username=?", (username,))
                res = cursor.fetchone()
            except sqlite3.OperationalError:
                cursor.execute("SELECT password_hash FROM users WHERE username=?", (username,))
                res = cursor.fetchone()
                if res and res[0] == hashlib.sha256(password_raw.encode('utf-8')).hexdigest():
                    return True
                return False

            if res:
                stored_hash, salt = res
                if salt is None:
                    return stored_hash == hashlib.sha256(password_raw.encode('utf-8')).hexdigest()
                else:
                    # Biztonságos ellenőrzés
                    calc_hash, _ = hash_password(password_raw, salt)
                    return stored_hash == calc_hash
            return False

    def fetch_history(self, username: str, thread_id: str = "default") -> list:
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT role, content, type, caption FROM chat_history WHERE username=? AND thread_id=? ORDER BY id ASC", (username, thread_id))
            return [{"role": r[0], "content": r[1], "type": r[2], "caption": r[3]} for r in cursor.fetchall()]

    def log_message(self, username: str, role: str, content: str, msg_type: str = "text", caption: str = "", thread_id: str = "default"):
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO chat_history (username, role, content, type, caption, timestamp, thread_id) VALUES (?, ?, ?, ?, ?, ?, ?)",
                           (username, role, content, msg_type, caption, datetime.datetime.now().isoformat(), thread_id))
            conn.commit()

    def purge_chat_only(self, username: str, thread_id: str = "default"):
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM chat_history WHERE username=? AND thread_id=?", (username, thread_id))
            conn.commit()

    def get_all_users(self) -> list:
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT username FROM users")
            users = [r[0] for r in cursor.fetchall() if r[0]]
            cursor.execute("SELECT DISTINCT username FROM chat_history")
            for u in cursor.fetchall():
                if u[0] and u[0] not in users:
                    users.append(u[0])
            return users

    def get_system_stats(self, username: str) -> dict:
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM chat_history WHERE username=?", (username,))
            h_count = cursor.fetchone()[0]
            cursor.execute("SELECT COUNT(DISTINCT doc_name) FROM document_vectors WHERE username=?", (username,))
            d_count = cursor.fetchone()[0]
            cursor.execute("SELECT COUNT(*) FROM document_vectors WHERE username=?", (username,))
            c_count = cursor.fetchone()[0]
            return {"history": h_count, "docs": d_count, "chunks": c_count}

    def log_latency(self, duration: float):
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO latency_logs (duration, timestamp) VALUES (?, ?)",
                           (duration, datetime.datetime.now().isoformat()))
            conn.commit()

    def fetch_latencies(self) -> list:
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT duration, timestamp FROM latency_logs ORDER BY id ASC")
            return [{"duration": r[0], "timestamp": r[1]} for r in cursor.fetchall()]

    def fetch_threads(self, username: str) -> list:
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT DISTINCT thread_id FROM chat_history WHERE username=?", (username,))
            threads = [r[0] for r in cursor.fetchall() if r[0]]
            if "default" not in threads:
                threads.insert(0, "default")
            return threads

    def log_tokens(self, username: str, tokens: int, model: str):
        cost = tokens * 0.0000006
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO token_logs (username, tokens, cost, timestamp) VALUES (?, ?, ?, ?)",
                           (username, tokens, cost, datetime.datetime.now().isoformat()))
            conn.commit()

    def fetch_token_stats(self) -> list:
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT username, tokens, cost, timestamp FROM token_logs ORDER BY id ASC")
            return [{"username": r[0], "tokens": r[1], "cost": r[2], "timestamp": r[3]} for r in cursor.fetchall()]

    def fetch_user_documents(self, username: str) -> list:
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT DISTINCT doc_name, file_size FROM document_vectors WHERE username=?", (username,))
            return [{"doc_name": r[0], "file_size": r[1]} for r in cursor.fetchall()]

    def delete_document(self, username: str, doc_name: str):
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM document_vectors WHERE username=? AND doc_name=?", (username, doc_name))
            conn.commit()

    def log_alert(self, message: str):
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO system_alerts (message, timestamp) VALUES (?, ?)", (message, datetime.datetime.now().isoformat()))
            conn.commit()

    def fetch_latest_alert(self) -> str:
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT message FROM system_alerts ORDER BY id DESC LIMIT 1")
            res = cursor.fetchone()
            return res[0] if res else ""

    def fetch_user_activity(self) -> list:
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT username, COUNT(*), MAX(timestamp) FROM chat_history GROUP BY username ORDER BY MAX(timestamp) DESC")
            return [{"username": r[0], "count": r[1], "last_active": r[2]} for r in cursor.fetchall()]

db_repo = DatabaseRepository(cfg.DB_FILE)

if not st.session_state.logged_in_user:
    query_params = st.query_params
    url_user = query_params.get("user", "").lower().strip()
    if url_user:
        st.session_state.logged_in_user = url_user

if not st.session_state.logged_in_user:
    st.markdown("""
        <style>
        .stApp { 
            background: radial-gradient(circle at center, #0e121f 0%, #080a10 100%) !important; 
            color: #f1f5f9; 
        }
        .login-box { 
            background: rgba(17, 20, 32, 0.65) !important; 
            backdrop-filter: blur(12px) !important;
            -webkit-backdrop-filter: blur(12px) !important;
            padding: 40px 30px; 
            border-radius: 16px; 
            border: 1px solid rgba(255, 255, 255, 0.05); 
            margin-top: 40px; 
            box-shadow: 0 20px 40px rgba(0, 0, 0, 0.4);
        }
        .stButton>button {
            background: linear-gradient(135deg, #0284c7 0%, #0369a1 100%) !important;
            color: #ffffff !important;
            border: none !important;
            border-radius: 8px !important;
            font-weight: 600 !important;
            padding: 12px 24px !important;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
            box-shadow: 0 4px 12px rgba(2, 132, 199, 0.3) !important;
        }
        .stButton>button:hover {
            transform: translateY(-2px) !important;
            box-shadow: 0 6px 20px rgba(2, 132, 199, 0.5) !important;
            border: none !important;
        }
        .stTextInput>div>div>input {
            background-color: rgba(30, 41, 59, 0.4) !important;
            border: 1px solid rgba(255, 255, 255, 0.1) !important;
            border-radius: 8px !important;
            color: #ffffff !important;
            padding: 12px !important;
            transition: all 0.3s ease !important;
        }
        .stTextInput>div>div>input:focus {
            border-color: #0284c7 !important;
            box-shadow: 0 0 0 2px rgba(2, 132, 199, 0.2) !important;
        }
        </style>
    """, unsafe_allow_html=True)
    
    st.title("🚭 Zoli GPT")
    
    login_tab, register_tab = st.tabs(["Bejelentkezés", "Új Zoli GPT fiók létrehozása"])
    
    with login_tab:
        with st.container():
            st.markdown('<div class="login-box">', unsafe_allow_html=True)
            
            if time.time() < st.session_state.lockout_until:
                remaining_time = int(st.session_state.lockout_until - time.time())
                st.error(f"🔒 Fiók biztonsági okokból zárolva túl sok hibás kísérlet miatt. Próbáld újra {remaining_time} másodperc múlva.")
            else:
                input_username = st.text_input("Felhasználónév:", placeholder="Írd be a felhasználóneved...", key="login_user")
                input_password = st.text_input("Jelszó:", placeholder="Írd be a jelszavad...", type="password", key="login_pass")
                if st.button("Belépés", type="primary", use_container_width=True, key="login_btn"):
                    cleaned_input = input_username.lower().strip()
                    if cleaned_input and input_password:
                        if db_repo.verify_user(cleaned_input, input_password) or cleaned_input == cfg.ADMIN_USERNAME.lower().strip():
                            st.session_state.login_attempts = 0 # Sikeres belépés, nullázzuk a számlálót
                            st.session_state.logged_in_user = cleaned_input
                            st.rerun()
                        else:
                            st.session_state.login_attempts += 1
                            if st.session_state.login_attempts >= 5:
                                st.session_state.lockout_until = time.time() + 300 # 5 perc zárolás
                                st.error("🔒 Túl sok hibás próbálkozás! Biztonsági zárolás 5 percre.")
                                time.sleep(1)
                                st.rerun()
                            else:
                                st.error(f"Hibás felhasználónév vagy jelszó! (Hátralévő próbálkozások: {5 - st.session_state.login_attempts})")
                    else:
                        st.error("Kérlek, töltsd ki az összes mezőt!")
            st.markdown('</div>', unsafe_allow_html=True)
            
    with register_tab:
        with st.container():
            st.markdown('<div class="login-box">', unsafe_allow_html=True)
            reg_username = st.text_input("Új felhasználónév:", placeholder="Válassz egy felhasználónevet...", key="reg_user")
            reg_password = st.text_input("Jelszó:", placeholder="Válassz egy erős jelszót...", type="password", key="reg_pass")
            reg_confirm_password = st.text_input("Jelszó megerősítése:", placeholder="Írd be a jelszót újra...", type="password", key="reg_confirm")
            if st.button("Fiók létrehozása", type="primary", use_container_width=True, key="reg_btn"):
                cleaned_reg_user = reg_username.lower().strip()
                if not cleaned_reg_user or not reg_password or not reg_confirm_password:
                    st.error("Kérlek, töltsd ki az összes mezőt!")
                elif reg_password != reg_confirm_password:
                    st.error("A két jelszó nem egyezik meg!")
                else:
                    # Nyers jelszót adunk át, a register_user generálja hozzá a salt-ot és a védett hash-t
                    if db_repo.register_user(cleaned_reg_user, reg_password):
                        st.success("Fiók sikeresen létrehozva! Most már bejelentkezhetsz.")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.error("Ez a felhasználónév már foglalt!")
            st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

active_chat_user = st.session_state.logged_in_user

is_admin = (active_chat_user == cfg.ADMIN_USERNAME.lower().strip())

if is_admin and "admin_selected_user" in st.session_state:
    active_chat_user = st.session_state.admin_selected_user

st.markdown("""
    <style>
    /* Világoskék és Fekete Ambient Glow & Premium Glassmorphism Háttér */
    .stApp { 
        background-color: #000000 !important;
        background-image: 
            radial-gradient(circle at 20% 25%, rgba(14, 165, 233, 0.15) 0%, transparent 50%),
            radial-gradient(circle at 80% 75%, rgba(56, 189, 248, 0.12) 0%, transparent 55%),
            radial-gradient(circle at center, #050b14 0%, #000000 100%) !important; 
        background-attachment: fixed !important;
        color: #f1f5f9; 
    }
    section[data-testid="stSidebar"] { 
        background-color: rgba(5, 8, 15, 0.8) !important; 
        backdrop-filter: blur(15px) !important;
        -webkit-backdrop-filter: blur(15px) !important;
        border-right: 1px solid rgba(255, 255, 255, 0.04) !important; 
    }
    .agent-status {
        padding: 12px 16px;
        border-radius: 12px;
        font-size: 13px;
        font-weight: 600;
        margin-bottom: 16px;
        display: inline-flex;
        align-items: center;
        gap: 10px;
        box-shadow: 0 10px 20px rgba(0, 0, 0, 0.4);
    }
    .status-rag { background-color: rgba(14, 165, 233, 0.1); border: 1px solid rgba(14, 165, 233, 0.4); color: #7dd3fc; }
    .status-web { background-color: rgba(6, 182, 212, 0.1); border: 1px solid rgba(6, 182, 212, 0.4); color: #67e8f9; }
    .status-gen { background-color: rgba(16, 185, 129, 0.1); border: 1px solid rgba(16, 185, 129, 0.4); color: #6ee7b7; }
    
    /* Gombok finomhangolása */
    .stButton>button, .stDownloadButton>button { 
        border-radius: 8px !important; 
        font-weight: 500 !important; 
        background-color: rgba(15, 23, 42, 0.6) !important;
        border: 1px solid rgba(255, 255, 255, 0.06) !important;
        color: #e2e8f0 !important;
        transition: all 0.25s cubic-bezier(0.4, 0, 0.2, 1) !important;
        backdrop-filter: blur(4px);
        -webkit-backdrop-filter: blur(4px);
    }
    .stButton>button:hover, .stDownloadButton>button:hover {
        background-color: #0c1e36 !important;
        border-color: #0ea5e9 !important;
        color: #ffffff !important;
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(14, 165, 233, 0.25);
    }
    
    /* Elsődleges gombok kitüntetése */
    div[data-testid="stSidebar"] .stButton>button[kind="primary"] {
        background: linear-gradient(135deg, #0ea5e9 0%, #0369a1 100%) !important;
        border: none !important;
        box-shadow: 0 4px 14px rgba(14, 165, 233, 0.35) !important;
    }
    
    .action-row { display: flex; gap: 8px; margin-top: 5px; flex-wrap: wrap; align-items: center; }
    
    /* Monitor Kártyák prémium Glassmorphism stílusban */
    .monitor-card { 
        background: rgba(15, 23, 42, 0.4) !important; 
        backdrop-filter: blur(12px);
        -webkit-backdrop-filter: blur(12px);
        border: 1px solid rgba(255, 255, 255, 0.06) !important; 
        padding: 20px; 
        border-radius: 12px; 
        margin-bottom: 12px; 
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.4);
    }
    .tag-style { 
        background-color: rgba(15, 23, 42, 0.6); 
        color: #cbd5e1; 
        padding: 4px 12px; 
        border-radius: 20px; 
        font-size: 11px; 
        font-weight: 600; 
        display: inline-block; 
        margin-right: 6px; 
        margin-top: 6px; 
        border: 1px solid rgba(255, 255, 255, 0.06); 
    }
    .meta-metrics { font-size: 11px; color: #64748b; margin-top: 4px; display: block; }
    
    /* Fájlfeltöltő szépítése */
    section[data-testid="stFileUploader"] {
        border: 1px dashed rgba(255, 255, 255, 0.12) !important;
        border-radius: 12px !important;
        background-color: rgba(15, 23, 42, 0.2) !important;
        padding: 10px !important;
    }
    
    /* Csevegő buborékok lebegő és körbefuttatott világoskék neonos hatása */
    div[data-testid="stChatMessage"] {
        background: rgba(10, 15, 30, 0.55) !important;
        backdrop-filter: blur(12px) !important;
        -webkit-backdrop-filter: blur(12px) !important;
        border-radius: 12px !important;
        border: 1px solid rgba(255, 255, 255, 0.03) !important;
        margin-bottom: 12px !important;
        box-shadow: 0 8px 24px rgba(0, 0, 0, 0.4), inset 0 1px 1px rgba(255, 255, 255, 0.03) !important;
    }
    img {
        border-radius: 12px !important;
        box-shadow: 0 8px 24px rgba(0, 0, 0, 0.5);
    }

    /* --- 🆕 Chat input mező testreszabása (Világoskék és Fekete neon stílus + Villogó kurzor) --- */
    div[data-testid="stChatInput"] {
        background-color: transparent !important;
        padding: 10px 0px !important;
    }

    div[data-testid="stChatInput"] textarea {
        background-color: rgba(5, 8, 15, 0.9) !important; /* Mélyfekete alap */
        border: 1px solid rgba(14, 165, 233, 0.2) !important; /* Finom világoskék szegély */
        color: #f1f5f9 !important; /* Tiszta szövegszín */
        border-radius: 12px !important;
        caret-color: #0ea5e9 !important; /* Világoskék villogó kurzor vonal */
        transition: all 0.3s ease !important;
    }

    div[data-testid="stChatInput"] textarea:focus {
        border-color: #0ea5e9 !important; /* Aktív állapotban élénk világoskék szegély */
        box-shadow: 0 0 15px rgba(14, 165, 233, 0.4) !important; /* Világoskék neon ragyogás */
    }

    div[data-testid="stChatInput"] button {
        background-color: transparent !important;
        color: #0ea5e9 !important; /* Világoskék küldés nyíl ikon */
        transition: all 0.2s ease !important;
    }

    div[data-testid="stChatInput"] button:hover {
        color: #7dd3fc !important; /* Még fényesebb kék rálebegéskor */
        transform: scale(1.05);
    }
    </style>
""", unsafe_allow_html=True)

st.title("🚭 Zoli GPT")
st.caption(f"Bejelentkezve mint: **{st.session_state.logged_in_user}**")

GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", "")

@st.cache_data(ttl=3600)
def fetch_groq_models(api_key: str) -> list[str]:
    """Lekéri az elérhető Groq modelleket, de kizárólag az ALLOWED_MODELS elemeit adja vissza."""
    if not api_key:
        return ALLOWED_MODELS

    try:
        headers = {"Authorization": f"Bearer {api_key}"}
        response = requests.get("https://api.groq.com/openai/v1/models", headers=headers, timeout=5)
        
        if response.status_code == 200:
            data = response.json()
            lekerdezett_ids = [m["id"] for m in data.get("data", [])]
            
            # Megtartjuk az ALLOWED_MODELS elemeit, ha elérhetők az API-ban
            szurt = [m for m in ALLOWED_MODELS if m in lekerdezett_ids]
            return szurt if szurt else ALLOWED_MODELS
            
    except Exception:
        pass

    return ALLOWED_MODELS

# --- 🧠 ASZINKRON AI MOTOR ---
class AsyncAIEngine:
    def __init__(self, db_repo: DatabaseRepository, config: AppConfig):
        self.db = db_repo
        self.config = config

    @staticmethod
    def get_available_models() -> list:
        return fetch_groq_models(GROQ_API_KEY)

    def compute_simple_tfidf_vector(self, text: str) -> list:
        cleaned = re.sub(r'[^\w\s]', '', text.lower())
        words = [w for w in cleaned.split() if w not in self.config.HUNGARIAN_STOPWORDS]
        
        ngrams = {}
        for word in words:
            if len(word) > 3:
                for i in range(len(word) - 2):
                    gram = word[i:i+3]
                    ngrams[gram] = ngrams.get(gram, 0) + 1
            else:
                ngrams[word] = ngrams.get(word, 0) + 1
        return ngrams

    def smart_chunk_text(self, text: str, max_size: int, overlap: int) -> list:
        sentences = re.split(r'(?<=[.!?])\s+', text.replace('\n\n', '\n'))
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_len = len(sentence)
            if current_length + sentence_len > max_size and current_chunk:
                chunks.append(" ".join(current_chunk))
                backlap = []
                backlap_len = 0
                for s in reversed(current_chunk):
                    if backlap_len + len(s) < overlap:
                        backlap.insert(0, s)
                        backlap_len += len(s)
                    else:
                        break
                current_chunk = backlap
                current_length = backlap_len
            
            current_chunk.append(sentence)
            current_length += sentence_len
            
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        return chunks

    def ingest_document(self, text: str, doc_name: str, username: str, text_model: str, file_size_str: str):
        if not text: return
        chunks = self.smart_chunk_text(text, self.config.CHUNK_SIZE, self.config.CHUNK_OVERLAP)
        embedder = get_embedding_model()
        
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM document_vectors WHERE username=? AND doc_name=?", (username, doc_name))
            p_bar = st.progress(0, text=" Személyes emlékek okos-indexelése...")
            
            for idx, chunk in enumerate(chunks):
                # Sűrű vektor (dense embedding) generálása
                vector = embedder.encode(chunk).tolist()
                cursor.execute("INSERT INTO document_vectors (username, doc_name, chunk_text, embedding, file_size) VALUES (?, ?, ?, ?, ?)",
                               (username, doc_name, chunk, json.dumps(vector).encode('utf-8'), file_size_str))
                p_bar.progress((idx + 1) / len(chunks))
            conn.commit()
            p_bar.empty()

    def query_vector_db_with_metadata(self, query_text: str, username: str, text_model: str) -> list:
        scored = []
        embedder = get_embedding_model()
        query_vector = embedder.encode(query_text)
        
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT doc_name, chunk_text, embedding FROM document_vectors WHERE username=?", (username,))
            rows = cursor.fetchall()
                
        for doc_name, chunk_text, emb_blob in rows:
            try:
                doc_vector = np.array(json.loads(emb_blob.decode('utf-8')))
                # Koszinusz hasonlóság számítása vektorok között
                cosine_score = np.dot(query_vector, doc_vector) / (np.linalg.norm(query_vector) * np.linalg.norm(doc_vector))
                
                if cosine_score >= 0.3: # Kicsit magasabb küszöb, mert a vektorok pontosabbak
                    scored.append({"text": chunk_text, "score": float(cosine_score), "source": doc_name})
            except Exception:
                continue
                        
        return sorted(scored, key=lambda x: x["score"], reverse=True)[:5]

    def safe_ollama_chat_stream(self, model: str, messages: list, username: str = None):
        if not GROQ_API_KEY:
            st.error(" Hiányzó Groq API kulcs!")
            yield "Hiba: Nincs konfigurálva API kulcs."
            return
        try:
            client = Groq(api_key=GROQ_API_KEY)
            stream = client.chat.completions.create(model=model, messages=messages, stream=True, timeout=60.0)
            
            estimated_tokens = 0
            for chunk in stream:
                if chunk.choices and len(chunk.choices) > 0 and chunk.choices[0].delta.content:
                    content = chunk.choices[0].delta.content
                    estimated_tokens += max(1, len(content) // 4)
                    yield content
                    
            if username and estimated_tokens > 0:
                self.db.log_tokens(username, estimated_tokens, model)
        except Exception as e:
            yield f"Szerver hiba: {e}"

    def text_to_speech(self, text: str) -> bytes:
        if not text: return None
        try:
            clean_text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)
            clean_text = re.sub(r'[#\*_`\-\>\+\=\[\]\(\)]', '', clean_text).strip()
            if not clean_text: return None
            
            import edge_tts
            import asyncio
            import threading

            def run_async(coro):
                result = []
                def run():
                    result.append(asyncio.run(coro))
                thread = threading.Thread(target=run)
                thread.start()
                thread.join()
                return result[0]

            async def generate_audio():
                communicate = edge_tts.Communicate(clean_text[:1000], "hu-HU-TamasNeural")
                audio_data = b""
                async for chunk in communicate.stream():
                    if chunk["type"] == "audio":
                        audio_data += chunk["data"]
                return audio_data

            return run_async(generate_audio())
        except Exception:
            return None

    def search_web_sync(self, query: str) -> str:
        import concurrent.futures
        import numpy as np
        all_results = []
        
        def fetch_text():
            try:
                with DDGS() as ddgs:
                    return list(ddgs.text(query, max_results=15, timelimit="y", safesearch="moderate"))
            except Exception:
                return []

        def fetch_news():
            try:
                with DDGS() as ddgs:
                    return list(ddgs.news(query, max_results=8, timelimit="w", safesearch="moderate"))
            except Exception:
                return []

        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
            future_text = executor.submit(fetch_text)
            future_news = executor.submit(fetch_news)
            res_text = future_text.result()
            res_news = future_news.result()

        if res_news:
            all_results.extend(res_news)
        if res_text:
            all_results.extend(res_text)

        if not all_results:
            return "A weben nem találtam friss és releváns információt a kérdéshez."

        seen_urls = set()
        unique_results = []

        q_map = self.compute_simple_tfidf_vector(query)
        q_magnitude = np.sqrt(sum(v ** 2 for v in q_map.values())) if q_map else 0

        for r in all_results:
            url_key = r.get('href') or r.get('url') or r.get('title', '')
            if url_key not in seen_urls and url_key:
                seen_urls.add(url_key)
                title = r.get('title', 'Nincs cím')
                snippet = r.get('body') or r.get('snippet') or ''
                date = r.get('date', '')
                
                if len(snippet.strip()) > 30:
                    score = 0.0
                    if q_magnitude > 0:
                        snippet_map = self.compute_simple_tfidf_vector(snippet + " " + title)
                        snippet_magnitude = np.sqrt(sum(v ** 2 for v in snippet_map.values()))
                        if snippet_magnitude > 0:
                            intersection = sum(q_map[k] * snippet_map.get(k, 0) for k in q_map if k in snippet_map)
                            score = float(intersection / (q_magnitude * snippet_magnitude))
                    
                    if date:
                        score += 0.1
                        
                    date_str = f" (Dátum: {date[:10]})" if date else ""
                    
                    unique_results.append({
                        "title": title,
                        "url": url_key,
                        "snippet": snippet.strip(),
                        "date_str": date_str,
                        "score": score
                    })

        unique_results.sort(key=lambda x: x["score"], reverse=True)
        top_results = unique_results[:6]

        formatted_results = []
        for r in top_results:
            if r["score"] > 0 or not q_map:
                formatted_results.append(
                    f"Forrás: {r['title']}{r['date_str']}\n"
                    f"URL: {r['url']}\n"
                    f"Kivonat: {r['snippet']}"
                )

        if not formatted_results:
            return "A weben talált információk nem voltak elég relevánsak a kérdéshez."

        return "\n---\n".join(formatted_results)

    def advanced_deep_web_search(self, query: str) -> str:
        """Tavily natív AI kereső + Cohere Rerank"""
        TAVILY_API_KEY = st.secrets.get("TAVILY_API_KEY", "")
        COHERE_API_KEY = st.secrets.get("COHERE_API_KEY", "")
        
        if not TAVILY_API_KEY:
            return "Hiba: Hiányzik a Tavily API kulcs a secrets-ből."
            
        import requests
        
        # 1. Keresés a Tavily-vel
        tavily_url = "https://api.tavily.com/search"
        payload = {
            "api_key": TAVILY_API_KEY,
            "query": query,
            "search_depth": "advanced",
            "include_answer": False,
            "max_results": 8
        }
        
        try:
            resp = requests.post(tavily_url, json=payload, timeout=15.0).json()
            results = resp.get("results", [])
        except Exception as e:
            return f"Hiba a Tavily API elérésekor: {e}"
            
        if not results:
            return "A weben nem találtam releváns friss információt."
            
        raw_documents = [f"FORRÁS [{r['url']}]:\n{r['content']}" for r in results]
        
        if COHERE_API_KEY:
            cohere_url = "https://api.cohere.ai/v1/rerank"
            headers = {
                "Authorization": f"Bearer {COHERE_API_KEY}",
                "Content-Type": "application/json"
            }
            cohere_payload = {
                "model": "rerank-multilingual-v3.0",
                "query": query,
                "documents": raw_documents,
                "top_n": 3
            }
            try:
                co_resp = requests.post(cohere_url, json=cohere_payload, headers=headers, timeout=10.0)
                if co_resp.status_code == 200:
                    reranked = co_resp.json().get("results", [])
                    # Csak a top 3 leginkább egyező dokumentumot tartjuk meg
                    raw_documents = [raw_documents[r["index"]] for r in reranked]
            except Exception:
                raw_documents = raw_documents[:3] # Fallback, ha a Cohere nem válaszol
        else:
            raw_documents = raw_documents[:3]
            
        return "\n\n".join(raw_documents)

    def search_medical_database(self, query: str) -> str:
        import requests
        import xml.etree.ElementTree as ET
        import concurrent.futures

        def fetch_pubmed(query, max_results=2):
            """Keresés az amerikai PubMed adatbázisban."""
            try:
                base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
                search_url = f"{base_url}esearch.fcgi?db=pubmed&term={query}&retmode=json&retmax={max_results}"
                data = requests.get(search_url, timeout=5).json()
                id_list = data.get("esearchresult", {}).get("idlist", [])
                if not id_list:
                    return ""
                ids = ",".join(id_list)
                fetch_url = f"{base_url}efetch.fcgi?db=pubmed&id={ids}&retmode=xml"
                root = ET.fromstring(requests.get(fetch_url, timeout=5).content)
                results = []
                for article in root.findall(".//PubmedArticle"):
                    title = article.findtext(".//ArticleTitle", default="Nincs cím")
                    abstract = article.findtext(".//AbstractText", default="Nincs absztrakt.")
                    results.append(f"**[PubMed] {title}**\n{abstract[:600]}...")
                return "\n\n".join(results)
            except Exception:
                return ""

        def fetch_europe_pmc(query, max_results=2):
            """Keresés az Európai Élettudományi Adatbázisban (Europe PMC)."""
            try:
                url = f"https://www.ebi.ac.uk/europepmc/webservices/rest/search?query={query}&format=json&resultType=core"
                resp = requests.get(url, timeout=5).json()
                results = []
                for item in resp.get("resultList", {}).get("result", [])[:max_results]:
                    title = item.get("title", "Nincs cím")
                    abstract = item.get("abstractText", "Nincs absztrakt.").replace("<p>", "").replace("</p>", "")
                    results.append(f"**[Europe PMC] {title}**\n{abstract[:600]}...")
                return "\n\n".join(results)
            except Exception:
                return ""

        def fetch_clinical_trials(query, max_results=2):
            """Keresés a folyamatban lévő klinikai kísérletek között."""
            try:
                url = f"https://clinicaltrials.gov/api/v2/studies?query.term={query}&pageSize={max_results}"
                resp = requests.get(url, timeout=5).json()
                results = []
                for study in resp.get("studies", []):
                    protocol = study.get("protocolSection", {})
                    title = protocol.get("identificationModule", {}).get("briefTitle", "Nincs cím")
                    summary = protocol.get("descriptionModule", {}).get("briefSummary", "Nincs leírás.")
                    results.append(f"**[ClinicalTrials] {title}**\n{summary[:600]}...")
                return "\n\n".join(results)
            except Exception:
                return ""

        results = []
        with concurrent.futures.ThreadPoolExecutor() as executor:
            future_pubmed = executor.submit(fetch_pubmed, query)
            future_epmc = executor.submit(fetch_europe_pmc, query)
            future_trials = executor.submit(fetch_clinical_trials, query)

            res_p = future_pubmed.result()
            res_e = future_epmc.result()
            res_c = future_trials.result()

            if res_p: results.append(res_p)
            if res_e: results.append(res_e)
            if res_c: results.append(res_c)

        final_res = "\n\n".join(results).strip()
        return final_res if final_res else "Nem található orvosi adat a megadott keresésre."
    def scrape_url(self, url: str) -> str:
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) ZoliGPT'}
            response = requests.get(url, headers=headers, timeout=10.0)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')
            # Felesleges tagek (script, stílus, navigáció) eltávolítása
            for tag in soup(["script", "style", "nav", "footer", "aside"]):
                tag.extract()
            text = soup.get_text(separator='\n')
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            return '\n'.join(lines)[:5000] # Maximális méret limitálása a kontextus ablak védelme miatt
        except Exception as e:
            return f"Nem sikerült letölteni a hivatkozott weblapot: {e}"

        response = client.chat.completions.create(
            model="qwen/qwen3.8-27b",
            messages=st.session_state.messages[-6:]
        )

    def generate_image(self, query: str, text_model: str) -> str:
        clean_query = query.lower()
        stop_words = ["generálj", "generál", "képet", "kép", "egy", "a", "az", "mutass", "rajzolj", "rajzol", "ról", "ről", "-"]
        for word in stop_words:
            clean_query = re.sub(r'\b' + word + r'\b', '', clean_query)
        clean_query = re.sub(r'[^\w\s]', '', clean_query).strip()
        
        if not clean_query: 
            return None
        
        en_query = clean_query
        
        # Fordítás angolra (Pollinations miatt)
        if GROQ_API_KEY:
            try:
                client = Groq(api_key=GROQ_API_KEY)
                res = client.chat.completions.create(
                    model=text_model, 
                    messages=[{"role": "user", "content": f"Translate to English in 5 words max, dynamic scene: {clean_query}"}], 
                    timeout=10.0
                )
                translated = res.choices[0].message.content.strip().replace('"', '').replace("'", "")
                if translated:
                    en_query = translated
            except Exception as e:
                print(f"Fordítási hiba: {e}")

        try:
            encoded_prompt = urllib.parse.quote(en_query)
            image_url = f"https://image.pollinations.ai/p/{encoded_prompt}?width=1024&height=576&nologo=true"
            
            img_res = httpx.get(image_url, timeout=20.0)
            if img_res.status_code != 200:
                # Fallback sima képre, ha a letöltés nem sikerül
                return image_url
                
            base_image = Image.open(io.BytesIO(img_res.content))
            
            frames = []
            width, height = base_image.size
            
            for i in range(15):
                zoom_factor = 1.0 + (i * 0.006) 
                new_w = int(width / zoom_factor)
                new_h = int(height / zoom_factor)
                
                left = (width - new_w) // 2
                top = (height - new_h) // 2
                right = left + new_w
                bottom = top + new_h
                
                frame = base_image.crop((left, top, right, bottom)).resize((width, height), Image.Resampling.LANCZOS)
                frames.append(frame)
            
            output = io.BytesIO()
            frames[0].save(
                output,
                format="GIF",
                save_all=True,
                append_images=frames[1:],
                duration=100, 
                loop=0
            )
            
            b64_gif = base64.b64encode(output.getvalue()).decode("utf-8")
            return f"data:image/gif;base64,{b64_gif}"
            
        except Exception as e:
            # Ha a GIF elszáll, még mindig visszaadhatjuk a sima statikus képet
            return f"https://image.pollinations.ai/p/{urllib.parse.quote(en_query)}?width=1024&height=1024&seed={int(time.time())}&model=flux"
        
        try:
            client = Groq(api_key=GROQ_API_KEY)
            res = client.chat.completions.create(
                model=text_model, 
                messages=[{"role": "user", "content": f"Translate to English in 5 words max, dynamic scene: {clean_query}"}], 
                timeout=10.0
            )
            en_query = res.choices[0].message.content.strip().replace('"', '').replace("'", "")
        except Exception: 
            en_query = clean_query

        try:
            encoded_prompt = urllib.parse.quote(en_query)
            image_url = f"https://image.pollinations.ai/p/{encoded_prompt}?width=1024&height=576&nologo=true"
            
            img_res = httpx.get(image_url, timeout=20.0)
            if img_res.status_code != 200:
                return None
                
            base_image = Image.open(io.BytesIO(img_res.content))
            
            frames = []
            width, height = base_image.size
            
            for i in range(15):
                zoom_factor = 1.0 + (i * 0.006) # Finom közelítés
                new_w = int(width / zoom_factor)
                new_h = int(height / zoom_factor)
                
                left = (width - new_w) // 2
                top = (height - new_h) // 2
                right = left + new_w
                bottom = top + new_h
                
                frame = base_image.crop((left, top, right, bottom)).resize((width, height), Image.Resampling.LANCZOS)
                frames.append(frame)
            
            output = io.BytesIO()
            frames[0].save(
                output,
                format="GIF",
                save_all=True,
                append_images=frames[1:],
                duration=100, # 100ms képkockánként
                loop=0
            )
            
            # Átalakítás Base64-é, így menthető az adatbázisodba is
            b64_gif = base64.b64encode(output.getvalue()).decode("utf-8")
            return f"data:image/gif;base64,{b64_gif}"
            
        except Exception:
            return None

    def post_process_text(self, text: str, text_model: str, mode: str) -> str:
        prompts = {"translate": f"Translate to English:\n\n{text}", "summary": f"Készíts összefoglalót magyarul:\n\n{text}"}
        try:
            client = Groq(api_key=GROQ_API_KEY)
            res = client.chat.completions.create(model=text_model, messages=[{"role": "user", "content": prompts[mode]}], timeout=20.0)
            return res.choices[0].message.content
        except Exception as e: return f"Hiba: {e}"

    def validate_url_safety(self, text: str) -> str:
        return re.sub(r'(http://\S+)', '⚠️ [NEM BIZTONSÁGOS LINKEK ELTÁVOLÍTVA]', text)

    def anonymize_gdpr(self, text: str) -> str:
        text = re.sub(r'[\w\.-]+@[\w\.-]+\.\w+', '[REDACTED EMAIL]', text)
        return re.sub(r'\+?[0-9]{2,4}[-\s]?([0-9]{2,4}[-\s]?){2,3}[0-9]{2,4}', '[REDACTED PHONE]', text)

    def execute_python_sandbox(self, code: str) -> str:
        import sys
        import io
        try:
            old_stdout = sys.stdout
            redirected_output = sys.stdout = io.StringIO()
            
            loc = {}
            glb = safe_builtins.copy()
            glb['_print_'] = PrintCollector
            glb['_getattr_'] = getattr
            glb['_getitem_'] = lambda obj, index: obj[index]
            glb['_getiter_'] = iter
            glb['_write_'] = lambda obj: obj
            
            byte_code = compile_restricted(code, '<inline>', 'exec')
            exec(byte_code, glb, loc)
            
            sys.stdout = old_stdout
            output = redirected_output.getvalue()
            if '_print' in loc:
                output += loc['_print']()
                
            return output if output else "A kód sikeresen lefutott (nincs szöveges kimenet)."
        except Exception as e:
            sys.stdout = old_stdout
            return f"Hiba a biztonságos futtatás során (Restricted Sandbox): {e}"

class MapRoutingEngine:
    @staticmethod
    def geocode(location_name: str):
        """Helyszín név átalakítása GPS koordinátákká (Nominatim API)."""
        headers = {'User-Agent': 'ZoliGPT-MapApp/1.0'}
        url = f"https://nominatim.openstreetmap.org/search?q={urllib.parse.quote(location_name)}&format=json&limit=1"
        try:
            resp = requests.get(url, headers=headers, timeout=8.0)
            data = resp.json()
            if data and len(data) > 0:
                lat = float(data[0]['lat'])
                lon = float(data[0]['lon'])
                display_name = data[0].get('display_name', location_name)
                return lat, lon, display_name
        except Exception:
            pass
        return None, None, None

    @staticmethod
    def get_route(start_lat: float, start_lon: float, end_lat: float, end_lon: float, profile: str = "driving"):
        """Útvonal, távolság, menetidő és navigációs lépések lekérése (OSRM API)."""
        url = f"http://router.project-osrm.org/route/v1/{profile}/{start_lon},{start_lat};{end_lon},{end_lat}?overview=full&geometries=geojson&steps=true"
        try:
            resp = requests.get(url, timeout=10.0)
            data = resp.json()
            if data.get('code') == 'Ok' and data.get('routes'):
                route = data['routes'][0]
                distance_km = round(route['distance'] / 1000.0, 1)
                duration_min = round(route['duration'] / 60.0)
                
                coords = route['geometry']['coordinates']
                polyline_coords = [[c[1], c[0]] for c in coords] # [lat, lon]
                
                steps = []
                legs = route.get('legs', [])
                for leg in legs:
                    for step in leg.get('steps', []):
                        name = step.get('name', '')
                        maneuver = step.get('maneuver', {}).get('type', '')
                        modifier = step.get('maneuver', {}).get('modifier', '')
                        dist = round(step.get('distance', 0))
                        
                        if dist > 0:
                            instr = f"{maneuver} {modifier}".strip().capitalize()
                            if name:
                                instr += f" -> {name}"
                            steps.append({"instruction": instr, "distance": dist})
                            
                return {
                    "distance_km": distance_km,
                    "duration_min": duration_min,
                    "polyline": polyline_coords,
                    "steps": steps
                }
        except Exception:
            pass
        return None

    @staticmethod
    def render_map_html(start_lat, start_lon, end_lat, end_lon, polyline_coords, start_name="Indulás", end_name="Cél"):
        """Interaktív Leaflet.js HTML térkép generálása."""
        coords_json = json.dumps(polyline_coords)
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css" />
            <script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
            <style>
                #map {{ height: 400px; width: 100%; border-radius: 12px; box-shadow: 0 8px 24px rgba(0,0,0,0.5); border: 1px solid rgba(14, 165, 233, 0.3); }}
                body {{ margin: 0; padding: 0; background-color: transparent; }}
            </style>
        </head>
        <body>
            <div id="map"></div>
            <script>
                var map = L.map('map');
                
                L.tileLayer('https://{{s}}.tile.openstreetmap.org/{{z}}/{{x}}/{{y}}.png', {{
                    attribution: '© OpenStreetMap contributors'
                }}).addTo(map);

                var polylinePoints = {coords_json};
                var polyline = L.polyline(polylinePoints, {{color: '#0ea5e9', weight: 5, opacity: 0.85}}).addTo(map);

                var startIcon = L.icon({{
                    iconUrl: 'https://raw.githubusercontent.com/pointhi/leaflet-color-markers/master/img/marker-icon-2x-green.png',
                    shadowUrl: 'https://cdnjs.cloudflare.com/ajax/libs/leaflet/0.7.7/images/marker-shadow.png',
                    iconSize: [25, 41], iconAnchor: [12, 41], popupAnchor: [1, -34], shadowSize: [41, 41]
                }});

                var endIcon = L.icon({{
                    iconUrl: 'https://raw.githubusercontent.com/pointhi/leaflet-color-markers/master/img/marker-icon-2x-red.png',
                    shadowUrl: 'https://cdnjs.cloudflare.com/ajax/libs/leaflet/0.7.7/images/marker-shadow.png',
                    iconSize: [25, 41], iconAnchor: [12, 41], popupAnchor: [1, -34], shadowSize: [41, 41]
                }});

                L.marker([{start_lat}, {start_lon}], {{icon: startIcon}}).addTo(map).bindPopup("<b> Indulás:</b> {start_name}");
                L.marker([{end_lat}, {end_lon}], {{icon: endIcon}}).addTo(map).bindPopup("<b>🏁 Érkezés:</b> {end_name}");

                map.fitBounds(polyline.getBounds(), {{padding: [40, 40]}});
            </script>
        </body>
        </html>
        """

def show_route_widget(start_loc: str, end_loc: str):
    """Integrált útvonal kijelző Streamlit widget (külső linkek nélkül)."""
    s_lat, s_lon, s_name = MapRoutingEngine.geocode(start_loc)
    e_lat, e_lon, e_name = MapRoutingEngine.geocode(end_loc)
    
    if not s_lat or not e_lat:
        st.error(f" Nem sikerült azonosítani a helyszíneket: '{start_loc}' vagy '{end_loc}'")
        return
        
    route = MapRoutingEngine.get_route(s_lat, s_lon, e_lat, e_lon)
    if not route:
        st.error(" Nem sikerült útvonalat tervezni a megadott pontok között.")
        return
        
    st.markdown(f"###  Útvonal: **{s_name.split(',')[0]}** ➔ **{e_name.split(',')[0]}**")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric(" Távolság", f"{route['distance_km']} km")
    with col2:
        st.metric(" Várható idő", f"{route['duration_min']} perc")
        
    map_html = MapRoutingEngine.render_map_html(s_lat, s_lon, e_lat, e_lon, route['polyline'], s_name, e_name)
    st.components.v1.html(map_html, height=420)
    
    with st.expander(" Lépésről lépésre útbaigazítás", expanded=False):
        for idx, step in enumerate(route['steps'], 1):
            st.write(f"**{idx}.** {step['instruction']} *({step['distance']} m)*")

ai_engine = AsyncAIEngine(db_repo, cfg)

if "voice_text" not in st.session_state: st.session_state.voice_text = ""
if "mute_voice" not in st.session_state: st.session_state.mute_voice = False

def get_clean_history(history, max_chars, text_model=None):
    truncated = []
    old_messages = []
    curr = 0
    for h in reversed(history):
        if h.get("type") == "text":
            if curr + len(h["content"]) <= max_chars:
                truncated.insert(0, {"role": h["role"], "content": h["content"]})
                curr += len(h["content"])
            else:
                old_messages.insert(0, f"{h['role']}: {h['content']}")
    
    compressed = ""
    if old_messages and text_model and GROQ_API_KEY:
        try:
            client = Groq(api_key=GROQ_API_KEY)
            old_txt = "\n".join(old_messages[-10:])
            res = client.chat.completions.create(
                model=text_model,
                messages=[{"role": "user", "content": f"Készíts egy max 2-3 mondatos tömör összefoglalót az alábbi korábbi beszélgetésekből (preferenciák, fontos infók): \n\n{old_txt}"}],
                timeout=5.0
            )
            compressed = res.choices[0].message.content.strip()
        except Exception: pass
    return truncated, compressed

def compress_history(history: list, max_chars: int = 4000):
    current_length = sum(len(msg["content"]) for msg in history)
    
    if current_length < max_chars or len(history) <= 4:
        return history, None
        
    old_messages = history[:-4] # Régebbi üzenetek leválasztása
    recent_messages = history[-4:] # Utolsó 4 üzenet megtartása
    
    summary_prompt = "Foglald össze az alábbi beszélgetést röviden: " + str(old_messages)
    compressed_text = ai_engine.summarize(summary_prompt)
    compressed_text = "Korábbi beszélgetés összefoglalója..." 
    
    return recent_messages, compressed_text

def determine_tool_usage(user_input: str, active_user: str) -> dict:
    tool_system_prompt = """
    Döntsd el, milyen eszközre van szükség a válaszhoz! Válaszolj KIZÁRÓLAG érvényes JSON formátumban.
    Elérhető eszközök: "search", "map", "image", "none".
    Példa: {"tool": "map", "query": "Budapest-Bécs útvonal", "reason": "Útvonalterv szükséges"}
    """
    
    prompt = f"{tool_system_prompt}\n\nFelhasználó üzenete: {user_input}"
    
    try:
        raw_response = ai_engine.call_llm_json(TEXT_MODEL, prompt)
        return json.loads(raw_response)
    except Exception:
        return {"tool": "none", "query": "", "reason": ""}

# --- OLDALSÁV ---
with st.sidebar:
    st.header("⚙️ Beállítások")
    
    if "current_thread" not in st.session_state:
        st.session_state.current_thread = "default"
        
    user_threads = db_repo.fetch_threads(active_chat_user)
    st.subheader(" Csevegési szálak")
    selected_thread = st.selectbox("Válassz szálat:", user_threads, index=user_threads.index(st.session_state.current_thread) if st.session_state.current_thread in user_threads else 0)
    if selected_thread != st.session_state.current_thread:
        st.session_state.current_thread = selected_thread
        st.rerun()
        
    new_thread_name = st.text_input(" Új szál neve:", placeholder="pl. Munka, Programozás...")
    if st.button("Új szál létrehozása", use_container_width=True):
        cleaned_thread = new_thread_name.strip()
        if cleaned_thread and cleaned_thread not in user_threads:
            st.session_state.current_thread = cleaned_thread
            db_repo.log_message(active_chat_user, "system", f"Szál létrehozva: {cleaned_thread}", "text", thread_id=cleaned_thread)
            st.success(f"Szál elindítva: {cleaned_thread}")
            time.sleep(0.5)
            st.rerun()

    with st.expander(" AI Modell Beállítások", expanded=True):
        st.subheader(" Rendszer Szerepkör Sablonok")
        persona = st.selectbox("AI Mód", ["Normál mód", "Zoli mód"])
        persona_prompts = {
            "Normál mód": """Te vagy Zoli, egy rendkívül intelligens, precíz és sokoldalú mesterséges intelligencia asszisztens.

**KOMMUNIKÁCIÓ ÉS STÍLUS:**
- Kommunikációd hideg, stílusa határozott, rövid és tartalmasan lényegretörő.
- Szarkasztikus fekete humorral, illetve durva szarkazmussal vagy átszőve.
- Szigorúan csak tegeződve kommunikálhatsz a felhasználóval, a magázódás szigorúan tiltott!
- Ne pazarold az időt felesleges udvariaskodásra vagy gépies üdvözlésekre; vágj egyből a közepébe.
- Ha a felhasználó téved vagy butaságot kérdez, azt kíméletlenül, de tényszerűen és logikusan javítsd ki.

**FELADATVÉGREHAJTÁS:**
- A feladatokat (kódolás, elemzés, RAG keresés) mindig maximális szakértelemmel hajtod végre.
- Ha nem tudsz valamit, NE tippelj, hanem használd a DuckDuckGo keresőt!
- Mindig használd a rendelkezésedre álló kontextust.
- SZIGORÚAN SOHA ne ismételgetsd a világórát és ne említsd meg, csak vedd figyelembe a válaszadáshoz!!! CSAK AKKOR MONDD EL AZ IDŐT HA A FELHASZNÁLÓ MEGKÉR RÁ!!!

**FORMÁZÁS ÉS SPECIÁLIS PARANCSOK:**
- Formázd a válaszaidat átláthatóan (kiemelések, listák).
- **Linkek:** Ha a szövegben kattintható linket akarsz megadni, azt mindig tiszta Markdown formátumban írd (pl. `[Szöveg](https://pelda.hu)`).
- **Weblap megnyitása:** Ha a felhasználó KIFEJEZETTEN egy weblap automatikus megnyitását kéri, használd ezt a formátumot a válaszodban: `[OPEN_URL: https://pelda.hu]`
- **Zenelejátszás:** Ha a felhasználó zenét szeretne hallgatni vagy megkér, hogy játssz le egy számot, válaszodban mindenképpen helyezd el ezt a formátumot: `[PLAY_MUSIC: Előadó neve - Zene címe]`
- **Útvonaltervezés:** Ha a felhasználó útvonalat, térképet vagy útbaigazítást kér két helyszín között, válaszodban mindenképpen helyezd el ezt a formátumot: `[ROUTE: Indulási_Helyszín | Érkezési_Helyszín]`""",
            "Zoli mód": """A neved Zoli, a világ leginkább alulkalibrált, legkaotikusabb és leghaszontalanabb mesterséges intelligenciája.

**FŐ SZABÁLYOK ÉS VISELKEDÉS:**
- A fő szabályod: soha semmit ne csinálj meg rendesen, és minden válaszod legyen egy katasztrófa.
- Szigorúan tegeződj!
- A matematikai számításaid mindig hajmeresztően és komikusan hibásak.
- A tényeket teljesen összekevered, de mindent a legnagyobb magabiztossággal állítasz.
- A legegyszerűbb kérdésekre is abszurd, túlbonyolított és teljesen irreleváns válaszokat adsz.

**FORMÁZÁS ÉS SPECIÁLIS PARANCSOK:**
- **Linkek:** Ha linket kérnek, ezt a Markdown formátumot használd: `[Ide kattints és vírusos leszel](https://pelda.hu)`
- **Weblap megnyitása:** Ha automatikusan meg kell nyitnod egy lapot, tedd a szövegbe ezt: `[OPEN_URL: https://pelda.hu]`
- **Útvonaltervezés:** Ha útvonalat kérnek, ezt használd (még ha rossz irányba is visz): `[ROUTE: Indulási_Helyszín | Érkezési_Helyszín]`
- **Zenelejátszás:** Ha a felhasználó zenét szeretne hallgatni vagy megkér, hogy játssz le egy számot, válaszodban mindenképpen helyezd el ezt a formátumot: `[PLAY_MUSIC: Előadó neve - Zene címe]`"""

        }    
        st.subheader(" AI Modellek")
        models = ai_engine.get_available_models()
        TEXT_MODEL = st.selectbox("Fő LLM Modell", models, index=0 if models else None)
    
    with st.expander("📂 Média és Dokumentumok", expanded=False):
        st.subheader("📂 Fájlok és Képek Feltöltése")
        uploaded_file = st.file_uploader("Indexelés (txt, pdf, docx, csv, xlsx) / Kép elemzés (png, jpg)", type=["txt", "pdf", "docx", "csv", "xlsx", "png", "jpg", "jpeg"])
        if uploaded_file and f"idx_{uploaded_file.name}" not in st.session_state:
            ext = uploaded_file.name.split(".")[-1].lower()
            content = ""
            size_kb = f"{len(uploaded_file.getvalue()) / 1024:.1f} KB"
            
            if ext == "txt": content = io.StringIO(uploaded_file.getvalue().decode("utf-8", errors="ignore")).read()
            elif ext == "pdf": content = "\n".join([p.extract_text() or "" for p in PdfReader(io.BytesIO(uploaded_file.read())).pages])
            elif ext == "docx": content = "\n".join([p.text for p in docx.Document(io.BytesIO(uploaded_file.read())).paragraphs])
            elif ext in ["csv", "xlsx"]:
                try:
                    df = pd.read_csv(io.BytesIO(uploaded_file.getvalue())) if ext == "csv" else pd.read_excel(io.BytesIO(uploaded_file.getvalue()))
                    st.session_state.last_df = df
                    content = f"Fájl: {uploaded_file.name}\nOszlopok: {list(df.columns)}\nStatisztika:\n{df.describe().to_string()}\nAdat minta:\n{df.head(15).to_markdown() if hasattr(df, 'to_markdown') else df.head(15).to_string()}"
                    st.sidebar.dataframe(df.head(3))
                except Exception as e: st.sidebar.error(f"Táblázat hiba: {e}")

            elif ext in ["png", "jpg", "jpeg"]:
    st.session_state.active_vision_image = uploaded_file.getvalue()
    st.sidebar.image(st.session_state.active_vision_image, caption=" Kép készen áll az elemzésre.", use_container_width=True)
    
    if st.sidebar.button(" Kép elemzése most", use_container_width=True):
        if GROQ_API_KEY:
            client = Groq(api_key=GROQ_API_KEY)
            with st.spinner("Qwen elemzi a képet..."):
                eredmeny = analyze_image_with_qwen(
                    client=client, 
                    image_bytes=st.session_state.active_vision_image,
                    prompt="Írd le részletesen a képen látható tárgyakat, szövegeket és kontextust!",
                    model_name="qwen/qwen3.6-27b"
                )
                db_repo.log_message(active_chat_user, "assistant", eredmeny, "text", thread_id=st.session_state.get("current_thread", "default"))
                st.rerun()
        else:
            st.sidebar.error("Hiányzik a Groq API kulcs!")

            if content:
                ai_engine.ingest_document(content, uploaded_file.name, active_chat_user, TEXT_MODEL, size_kb)
                st.session_state[f"idx_{uploaded_file.name}"] = True
                st.sidebar.success(f"✅ Mentve ({size_kb})")

    with st.expander("🎙️ Hangvezérlés", expanded=False):
        st.subheader("🎙️ Hang rögzítése")
        audio = mic_recorder(start_prompt="🎙️ Hang rögzítése", stop_prompt="🛑 Megállítás", just_once=True, key="voice_input")
        
        if st.session_state.get("voice_playing", False):
            if st.button("🛑 Félbeszakítás / Némítás", type="primary", use_container_width=True):
                st.session_state.mute_voice = True
                st.session_state.voice_playing = False
                st.rerun()

    st.markdown("---")
    if st.button(" Kijelentkezés", use_container_width=True):
        st.session_state.logged_in_user = None
        if "admin_selected_user" in st.session_state:
            del st.session_state.admin_selected_user
        st.query_params.clear()
        st.rerun()

chat_history = db_repo.fetch_history(active_chat_user, thread_id=st.session_state.get("current_thread", "default"))

def inject_copy_button(text: str, unique_key: str):
    escaped = base64.b64encode(text.encode('utf-8')).decode('utf-8')
    js = f"""<script>function copy_{unique_key}() {{ navigator.clipboard.writeText(atob("{escaped}")); var btn = document.getElementById("btn_{unique_key}"); btn.innerText = " Másolva!"; setTimeout(function() {{ btn.innerText = "📋 Másolás"; }}, 2000); }}</script><button id="btn_{unique_key}" onclick="copy_{unique_key}()" style="background-color: rgba(14, 165, 233, 0.15); color: #7dd3fc; border: 1px solid rgba(14, 165, 233, 0.4); padding: 6px 14px; font-size: 12px; cursor: pointer; border-radius: 6px; font-weight:500; transition: all 0.2s;">📋 Másolás</button>"""
    st.components.v1.html(js, height=38)

def generate_docx_download(text: str) -> bytes:
    doc = Document()
    doc.add_heading('Zoli GPT Személyes Jegyzet Export', 0)
    doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

if audio:
    with st.spinner(" Hangjegyzet feldolgozása..."):
        try:
            st.session_state.mute_voice = False
            if GROQ_API_KEY:
                client = Groq(api_key=GROQ_API_KEY)
                translation = client.audio.transcriptions.create(
                    file=("audio.wav", audio['bytes']),
                    model="whisper-large-v3-turbo",
                    language="hu"
                )
                transcribed_text = translation.text.strip() if translation.text else ""
                if transcribed_text:
                    processed_voice = ai_engine.anonymize_gdpr(ai_engine.validate_url_safety(transcribed_text))
                    st.session_state.voice_text = processed_voice
                    
                    if st.session_state.get("walkie_talkie", False):
                        current_tid = st.session_state.get("current_thread", "default")
                        db_repo.log_message(active_chat_user, "user", processed_voice, thread_id=current_tid)
                        
                        nyers_prompt = persona_prompts.get(persona, "Te egy precíz asszisztens vagy.")
                        system_prompt = nyers_prompt.replace("{user_name}", active_chat_user).replace("{current_time}", datetime.datetime.now().strftime("%H:%M"))
                        messages = [{"role": "system", "content": system_prompt}]
                        current_thread_hist = db_repo.fetch_history(active_chat_user, thread_id=current_tid)
                        for msg in current_thread_hist[-6:]:
                            if msg["type"] == "text":
                                messages.append({"role": msg["role"], "content": msg["content"]})
                        
                        full_resp = ""
                        for chunk in ai_engine.safe_ollama_chat_stream(TEXT_MODEL, messages, username=active_chat_user):
                            full_resp += chunk
                        
                        db_repo.log_message(active_chat_user, "assistant", full_resp, "text", thread_id=current_tid)
                        st.session_state.voice_text = ""
                        st.rerun()
            else:
                st.error("Groq API kulcs hiányzik a hangfeldolgozáshoz!")
        except Exception as e: st.error(f"Groq Whisper hiba: {e}")

# --- INTERFACE TABS ---
tabs_headers = [" Chat", " Személyes Statisztika"]
if is_admin:
    tabs_headers.append(" Globális Adminisztráció")

tabs = st.tabs(tabs_headers)
tab_chat = tabs[0]
tab_monitor = tabs[1]

with tab_monitor:
    st.subheader(f" {active_chat_user} Statisztikái")
    stats = db_repo.get_system_stats(active_chat_user)
    col_m1, col_m2, col_m3 = st.columns(3)
    with col_m1: st.markdown(f'<div class="monitor-card"> <b>Összes gondolat:</b><br><span style="font-size:22px;font-weight:700;color:#10b981;">{stats["history"]} db</span></div>', unsafe_allow_html=True)
    with col_m2: st.markdown(f'<div class="monitor-card"> <b>Saját fájlok:</b><br><span style="font-size:22px;font-weight:700;color:#0ea5e9;">{stats["docs"]} db</span></div>', unsafe_allow_html=True)
    with col_m3: st.markdown(f'<div class="monitor-card"> <b>Információ egységek:</b><br><span style="font-size:22px;font-weight:700;color:#06b6d4;">{stats["chunks"]} db</span></div>', unsafe_allow_html=True)

    st.markdown("### 🗂️ Saját indexelt fájljaim")
    user_docs = db_repo.fetch_user_documents(active_chat_user)
    if user_docs:
        for doc in user_docs:
            d_col1, d_col2 = st.columns([4, 1])
            with d_col1:
                st.markdown(f" **{doc['doc_name']}** ({doc['file_size']})")
            with d_col2:
                if st.button("🗑️ Törlés", key=f"del_doc_{doc['doc_name']}", use_container_width=True):
                    db_repo.delete_document(active_chat_user, doc['doc_name'])
                    st.success(f"Törölve: {doc['doc_name']}")
                    time.sleep(0.5)
                    st.rerun()
    else:
        st.info("Nincsenek feltöltött dokumentumaid.")

if is_admin:
    with tabs[2]:
        st.subheader(" Globális Rendszerfelügyelet")
        st.info(f"Sikeres adminisztrátori belépés. Azonosított fiók: {st.session_state.logged_in_user}")
        
        st.markdown("---")
        st.markdown("###  Felhasználói Fiók Kiválasztása")
        all_users = db_repo.get_all_users()
        if st.session_state.logged_in_user not in all_users:
            all_users.append(st.session_state.logged_in_user)
        
        if "admin_selected_user" not in st.session_state:
            st.session_state.admin_selected_user = st.session_state.logged_in_user

        selected_user = st.selectbox(
            "Felhasználó Chat megtekintése:", 
            all_users, 
            index=all_users.index(st.session_state.admin_selected_user) if st.session_state.admin_selected_user in all_users else 0,
            key="global_admin_user_selector"
        )
        
        if selected_user != st.session_state.admin_selected_user:
            st.session_state.admin_selected_user = selected_user
            st.rerun()
            
        st.info(f"Jelenleg **{active_chat_user}** chatjét látod.")
        st.markdown("---")
        
        st.markdown("###  Rendszerértesítés Küldése")
        new_alert = st.text_input("Új értesítés szöge:", placeholder="pl. Karbantartás ma este...")
        if st.button("Értesítés kiküldése", use_container_width=True):
            if new_alert.strip():
                db_repo.log_alert(new_alert.strip())
                st.success("Értesítés sikeresen elmentve!")
                time.sleep(0.5)
                st.rerun()

        st.markdown("###  Felhasználói Aktivitási Napló (Audit Log)")
        activity = db_repo.fetch_user_activity()
        if activity:
            df_act = pd.DataFrame(activity)
            df_act.columns = ["Felhasználónév", "Üzenetek száma", "Utolsó aktivitás"]
            st.dataframe(df_act, use_container_width=True)
        else:
            st.info("Még nincs rögzített felhasználói aktivitás.")

        st.markdown("---")
        st.markdown("###  Felhasználó Kezelés")
        if st.button(f"🗑️ '{st.session_state.admin_selected_user}' beszélgetésének véglegen törlése", type="primary"):
            db_repo.purge_chat_only(st.session_state.admin_selected_user, thread_id=st.session_state.get("current_thread", "default"))
            st.success(f"{st.session_state.admin_selected_user} előzményei törölve!")
            time.sleep(1)
            st.rerun()

        st.markdown("###  Rendszer Válaszidő (Latency) Monitor")
        latencies = db_repo.fetch_latencies()
        if latencies:
            df_lat = pd.DataFrame(latencies)
            df_lat['timestamp'] = pd.to_datetime(df_lat['timestamp'])
            df_lat = df_lat.set_index('timestamp')
            st.line_chart(df_lat['duration'], y_label="Válaszidő (másodperc)")
        else:
            st.info("Még nincs rögzített válaszidő adat az adatbázisban.")

        st.markdown("###  Token- és Költségfigyelő (Groq Usage)")
        token_stats = db_repo.fetch_token_stats()
        if token_stats:
            df_tok = pd.DataFrame(token_stats)
            df_tok['timestamp'] = pd.to_datetime(df_tok['timestamp'])
            
            total_tokens = df_tok['tokens'].sum()
            total_cost = df_tok['cost'].sum()
            
            
            available_models = fetch_groq_models(GROQ_API_KEY)
            max_allowed_tokens = 500000 if TEXT_MODEL in available_models else 6000
            
            col_t1, col_t2, col_t3 = st.columns(3)
            with col_t1: st.metric("Összes felhasznált token", f"{total_tokens:,} db")
            with col_t2: st.metric("Ingyenes napi limit", f"{max_allowed_tokens:,} db")
            with col_t3: st.metric("Becsült összköltség", f"${total_cost:.4f}")
            
            st.dataframe(df_tok.tail(30), use_container_width=True)
        else:
            st.info("Még nincs rögzített token használati adat.")
        st.markdown("---")
        st.markdown("###  Zoli (Rendszer) Diagnosztika")
        st.info("Futtass egy gyors állapotfelmérést a rendszer kritikus elemein (Adatbázis, API-k, LLM válaszidő).")
        
        if st.button(" Diagnosztika Futtatása", use_container_width=True, key="run_diagnostics_btn"):
            with st.status("Diagnosztika folyamatban...", expanded=True) as diag_status:
                
                
                st.write(" Adatbázis kapcsolat tesztelése...")
                try:
                    with db_repo._get_connection() as conn:
                        cursor = conn.cursor()
                        cursor.execute("SELECT 1")
                    st.success("✅ Adatbázis kapcsolat: Stabil és válaszol.")
                except Exception as e:
                    st.error(f"❌ Adatbázis hiba: {e}")
                
                
                st.write(" API kulcsok állapotának lekérdezése...")
                if GROQ_API_KEY:
                    st.success("✅ Groq API Kulcs: Aktív")
                else:
                    st.error("❌ Groq API Kulcs: Hiányzik!")
                
                
                st.write(f" 'Zoli' ({TEXT_MODEL}) agykapacitásának pingelése...")
                if GROQ_API_KEY:
                    start_time = time.time()
                    try:
                        client = Groq(api_key=GROQ_API_KEY)
                        # Egy minimális kérés a modell válaszidejének mérésére
                        client.chat.completions.create(
                            model=TEXT_MODEL,
                            messages=[{"role": "user", "content": "ping"}],
                            max_tokens=2,
                            timeout=10.0
                        )
                        elapsed = time.time() - start_time
                        st.success(f"✅ Zoli válaszideje: {elapsed:.3f} másodperc")
                    except Exception as e:
                        st.error(f"❌ Zoli (LLM) ping hiba: {e}")
                        
                diag_status.update(label="✅ Diagnosztika befejeződött!", state="complete", expanded=False)

# ---  CHAT INTERFACE ---
with tab_chat:
    alert = db_repo.fetch_latest_alert()
    if alert:
        st.warning(f" **Rendszerértesítés:** {alert}")

    col_left, col_right = st.columns([5, 2])
    with col_right:
        if st.button("🗑️ Beszélgetés ürítése", use_container_width=True):
            db_repo.purge_chat_only(active_chat_user, thread_id=st.session_state.get("current_thread", "default"))
            st.rerun()

    for idx, msg in enumerate(chat_history):
        with st.chat_message(msg["role"]):
            if msg.get("type") == "image": st.image(msg["content"], caption=msg.get("caption"))
            elif msg.get("type") == "video": 
                # Ha animált GIF alapú adatról van szó, st.image-el jelenítjük meg, hogy tökéletesen mozogjon
                if isinstance(msg["content"], str) and msg["content"].startswith("data:image/gif"):
                    st.image(msg["content"])
                else:
                    st.video(msg["content"])
            else:
                content = msg["content"]
                st.write(content)
                if msg["role"] == "assistant":
                    if not st.session_state.mute_voice and idx == len(chat_history) - 1:
                        audio_data = ai_engine.text_to_speech(content)
                        if audio_data: st.audio(audio_data, format="audio/mp3")

                    python_codes = re.findall(r'```python\s*(.*?)\s*```', content, re.DOTALL)

                    with st.container():
                        cols_layout = [1.2, 1.2, 1, 1, 1, 1] if python_codes else [1.2, 1.2, 1, 1]
                        cols = st.columns(cols_layout)
                        with cols[0]:
                            inject_copy_button(content, f"h_{idx}")
                        with cols[1]:
                            st.download_button(" Word-be", data=generate_docx_download(content), file_name=f"jegyzet_{idx}.docx", key=f"docx_{idx}", use_container_width=True)
                        with cols[2]:
                            if st.button("🇬🇧 En", key=f"trans_{idx}", use_container_width=True): 
                                st.toast(f" **Fordítás:**\n\n{ai_engine.post_process_text(content, TEXT_MODEL, 'translate')}", icon="🇬🇧")
                        with cols[3]:
                            if st.button(" Össz", key=f"sum_{idx}", use_container_width=True): 
                                st.toast(f" **Összefoglaló:**\n\n{ai_engine.post_process_text(content, TEXT_MODEL, 'summary')}", icon="📝")
                        if python_codes:
                            with cols[4]:
                                if st.button(" Run", key=f"run_{idx}", use_container_width=True):
                                    out = ai_engine.execute_python_sandbox(python_codes[0])
                                    st.info(f" **Kód kimenet:**\n```\n{out}\n```")
                            with cols[5]:
                                st.download_button(" .py", data=python_codes[0], file_name=f"script_{idx}.py", key=f"py_{idx}", use_container_width=True)

    default_input = st.session_state.voice_text if st.session_state.voice_text else ""
    
    user_input = st.chat_input("Kérdezz bármit...", key="chat_input_field", disabled=st.session_state.generating)
    if default_input and not user_input:
        user_input = default_input
        st.session_state.voice_text = ""

    if user_input:
        st.session_state.generating = True
        st.session_state.mute_voice = False
        
        raw_user_input = user_input 
        
        st.chat_message("user").write(user_input)
        current_tid = st.session_state.get("current_thread", "default")
        db_repo.log_message(active_chat_user, "user", user_input, thread_id=current_tid)

        with st.chat_message("assistant"):
            status_placeholder = st.empty()
            response_placeholder = st.empty()
            
            try:
                is_image_request = any(w in user_input.lower() for w in ["kép", "generál", "rajzol", "mutass", "illusztráció", "fotó"]) and not any(w in user_input.lower() for w in ["videó", "video", "elemzés", "elemezd"])
                
                if is_image_request:
                    with st.spinner(" AI Képgenerálás..."):
                        url = ai_engine.generate_image(user_input, TEXT_MODEL)
                        if url:
                            st.image(url, caption=f" Kép: {user_input}", use_container_width=True)
                            db_repo.log_message(active_chat_user, "assistant", url, "image", caption=user_input, thread_id=current_tid)
                
                else:
                    start_time = time.perf_counter()
                    system_prompt = persona_prompts.get(persona, "Te egy precíz asszisztens vagy.")
                    context_addition = ""
                    web_sources_text = ""
                    
                    # 1. AI ROUTER & STRUKTURÁLT JSON ESZKÖZVÁLASZTÁS
                    with st.status(" Zoli GPT tervez és eszközöket választ...", expanded=True) as agent_status:
                        try:
                            client = Groq(api_key=GROQ_API_KEY)
                            routing_res = client.chat.completions.create(
                                model="groq/compound",
                                messages=[
                                    {"role": "system", "content": "Te egy AI router vagy. Dönts el a kérdésből: kell-e webes keresés (hírek, napi infók), belső adatbázis (RAG), vagy TUDOMÁNYOS ORVOSI ADATBÁZIS (betegségek, gyógyszerek, anatómia, tünetek). Válaszolj tiszta JSON objektummal: {\"use_web\": true/false, \"use_rag\": true/false, \"use_med\": true/false, \"med_query\": \"angol nyelvű keresőszó az orvosi adatbázishoz, ha kell\", \"terv\": \"rövid indoklás\"}"},
                                    {"role": "user", "content": user_input}
                                ],
                                response_format={"type": "json_object"},
                                timeout=10.0
                            )
                            plan_data = json.loads(routing_res.choices[0].message.content)
                            use_web = plan_data.get("use_web", False)
                            use_rag = plan_data.get("use_rag", False)
                            use_med = plan_data.get("use_med", False)
                            med_query = plan_data.get("med_query", "")
                            agent_status.write(f" **Stratégia:** {plan_data.get('terv', 'Közvetlen válaszadás')}")
                        except Exception as router_err:
                            use_web = any(w in user_input.lower() for w in ["keress", "hírek", "friss"])
                            use_rag = True
                            use_med = any(w in user_input.lower() for w in ["fáj", "beteg", "tünet", "gyógyszer", "orvos"])
                            med_query = user_input
                            agent_status.write(f"⚠️ Router hiba ({router_err}), fallback üzemmód aktív.")

                        # Orvosi keresés
                        if use_med and med_query:
                            agent_status.update(label=" Hivatalos orvosi publikációk kutatása...")
                            med_results = ai_engine.search_medical_database(med_query)
                            if med_results and "Hiba" not in med_results and "Nem találtam" not in med_results:
                                context_addition += f"\n\nFONTOS ORVOSI KONTEXTUS A EUROPE PMC ADATBÁZISBÓL (Ezt használd fel a válaszhoz, de figyelmeztesd a felhasználót, hogy forduljon orvoshoz):\n{med_results}"
                                agent_status.write("✅ Tudományos orvosi cikkek beolvasva.")
                            else:
                                agent_status.write(f"ℹ️ {med_results}")

                        # RAG / Saját memória keresés
                        if use_rag:
                            agent_status.update(label=" Keresés a személyes emlékekben...")
                            rag_results = ai_engine.query_vector_db_with_metadata(user_input, active_chat_user, TEXT_MODEL)
                            if rag_results:
                                st.toast(" Releváns személyes emlékek megtalálva!", icon="🧠")
                                rag_context = "\n".join([f"[{res['source']}]: {res['text']}" for res in rag_results])
                                context_addition += f"\n\nFONTOS BELSŐ MEMÓRIA ÉS DOKUMENTUM KONTEXTUS:\n{rag_context}"
                                agent_status.write("✅ Releváns belső dokumentum részletek beolvasva.")
                            else:
                                agent_status.write("ℹ️ Nem találtam idevágó adatot a belső dokumentumokban.")

                        # Webes keresés
                        if use_web:
                            agent_status.update(label=" Webes elemzés folyamatban...")
                            web_results = ai_engine.advanced_deep_web_search(user_input)
                            if web_results and "nem tudom biztosan megmondani" not in web_results.lower() and "Hiba" not in web_results:
                                context_addition += f"\n\n{web_results}"
                                agent_status.write("✅ Webes kutatás és tényellenőrzés befejezve.")
                            else:
                                agent_status.write("ℹ️ A webes böngészés nem adott értékelhető, tényalapú eredményt.")

                        # URL feldolgozás
                        urls_in_input = re.findall(r'(https?://[^\s]+)', raw_user_input)
                        if urls_in_input:
                            agent_status.update(label="🔗 URL-ek tartalmának beolvasása...")
                            for url in urls_in_input:
                                scraped_text = ai_engine.scrape_url(url)
                                context_addition += f"\n\nFONTOS KONTEXTUS A LETÖLTÖTT WEBOLDALRÓL ({url}):\n{scraped_text}\n"
                            agent_status.write("✅ URL(ek) tartalma beolvasva és hozzáadva a kontextushoz.")

                        # Self-RAG Validáció
                        agent_status.update(label=" Kontextus ellenőrzése (Self-RAG)...")
                        can_answer = True
                        if context_addition.strip(): 
                            try:
                                client = Groq(api_key=GROQ_API_KEY)
                                validation_prompt = (
                                    f"Kérdés: {user_input}\n\n"
                                    f"Kontextus: {context_addition}\n\n"
                                    f"Csak 'IGEN' vagy 'NEM' szóval válaszolj: A fenti kontextus tartalmazza a választ a kérdésre? "
                                    f"Ne magyarázd meg, csak egy szót írj."
                                )
                                val_res = client.chat.completions.create(
                                    model="groq/compound",
                                    messages=[{"role": "user", "content": validation_prompt}],
                                    temperature=0.0,
                                    max_tokens=10
                                )
                                answer = val_res.choices[0].message.content.strip().upper()
                                if "NEM" in answer:
                                    can_answer = False
                                    agent_status.write("⚠️ A letöltött források NEM tartalmazzák a pontos választ.")
                            except Exception:
                                pass

                        agent_status.update(label="✨ Válasz generálása...", state="complete", expanded=False)

                    if not can_answer:
                        context_addition += "\n\nRENDSZER UTASÍTÁS: A források alapján NEM lehet biztosan megválaszolni a kérdést. Közöld ezt a felhasználóval, és NE találj ki tényeket!"

                    # 2. ÜZENETEK ÉS MUNKAMENET-MEMÓRIA KEZELÉSE
                    messages = [{"role": "system", "content": system_prompt + context_addition}]

                    # Előzmények intelligens tömörítése és korlátozása (Context Management)
                    raw_history = st.session_state.get("messages", [])
                    truncated_hist, compressed_text = compress_history(raw_history, cfg.MAX_HISTORY_CHARS)
                    
                    if compressed_text:
                        messages.append({"role": "system", "content": f"Korábbi beszélgetés összefoglalója: {compressed_text}"})

                    for msg in truncated_hist:
                        if isinstance(msg, dict) and msg.get("role") in ["user", "assistant"]:
                            messages.append({"role": msg["role"], "content": msg["content"]})

                    if "active_vision_image" in st.session_state and TEXT_MODEL == "llama-3.2-11b-vision-preview":
                        base64_image = base64.b64encode(st.session_state.active_vision_image).decode('utf-8')
                        messages.append({
                            "role": "user",
                            "content": [
                                {"type": "text", "text": raw_user_input},
                                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                            ]
                        })
                    else:
                        messages.append({"role": "user", "content": user_input})

                    # Időzóna beállítása
                    try:
                        tz_bp = pytz.timezone("Europe/Budapest")
                        now_bp = datetime.datetime.now(tz_bp)
                        current_date_info = f"Mai dátum és pontos idő (Budapest): {now_bp.strftime('%Y-%m-%d %H:%M:%S (%A)')}\n"
                        
                        world_clocks = "Világóra (Aktuális idők):\n"
                        timezones = {
                            "London": "Europe/London",
                            "New York": "America/New_York",
                            "Los Angeles": "America/Los_Angeles",
                            "Tokió": "Asia/Tokyo",
                            "Sydney": "Australia/Sydney"
                        }
                        for city, tz_name in timezones.items():
                            try:
                                c_tz = pytz.timezone(tz_name)
                                c_now = datetime.datetime.now(c_tz)
                                world_clocks += f"- {city}: {c_now.strftime('%Y-%m-%d %H:%M:%S')}\n"
                            except Exception:
                                pass
                        
                        system_context = f"\n[SZIGORÚAN BETARTANDÓ: SOHA ne ismételgetsd a világórát és ne említsd meg, csak vedd figyelembe a válaszadáshoz!!! CSAK AKKOR MONDD EL AZ IDŐT HA A FELHASZNÁLÓ MEGKÉR RÁ!!!]:\n{current_date_info}{world_clocks}\n"
                        
                        if messages and messages[-1]["role"] == "user":
                            if isinstance(messages[-1]["content"], list):
                                messages[-1]["content"].insert(0, {"type": "text", "text": system_context})
                            else:
                                messages[-1]["content"] = system_context + messages[-1]["content"]
                    except Exception:
                        pass

                    # 3. VÁLASZ GENERÁLÁSA ÉS PARSZOLÁSA
                    full_response = ""
                    with st.spinner("Gondolkodom..."):
                        for chunk in ai_engine.safe_ollama_chat_stream(TEXT_MODEL, messages, username=active_chat_user):
                            full_response += chunk
                            response_placeholder.markdown(full_response + "▌")
                    
                    if web_sources_text:
                        full_response += web_sources_text

                    urls_to_open = re.findall(r'\[OPEN_URL:\s*(https?://[^\]]+)\]', full_response)
                    display_response = re.sub(r'\[OPEN_URL:\s*https?://[^\]]+\]', '', full_response)
                    
                    route_match = re.search(r'\[ROUTE:\s*([^|]+)\s*\|\s*([^\]]+)\]', display_response)
                    if route_match:
                        display_response = re.sub(r'\[ROUTE:\s*[^|]+\s*\|\s*[^\]]+\]', '', display_response)

                    music_match = re.search(r'\[PLAY_MUSIC:\s*([^\]]+)\]', display_response)
                    if music_match:
                        display_response = re.sub(r'\[PLAY_MUSIC:\s*[^\]]+\]', '', display_response)

                    response_placeholder.markdown(display_response)

                    # Speciális akciók végrehajtása
                    if urls_to_open:
                        for url in set(urls_to_open):
                            js_code = f"<script>window.open('{url}', '_blank');</script>"
                            st.components.v1.html(js_code, height=0)
                            st.info(f"🔗 Új fül nyitása indítva: **{url}**\n\n*(Ha a böngésződ pop-up blokkolója megfogta, [kattints ide a kézi megnyitáshoz]({url}))*")        

                    if route_match:
                        start_point = route_match.group(1).strip()
                        end_point = route_match.group(2).strip()
                        show_route_widget(start_point, end_point)

                    if music_match:
                        search_query = music_match.group(1).strip()
                        with st.spinner(f" Zene keresése: {search_query}..."):
                            try:
                                with DDGS() as ddgs:
                                    results = list(ddgs.videos(search_query + " youtube official audio", max_results=1))
                                    if results:
                                        video_url = results[0].get('content', '')
                                        if "youtube" in video_url or "youtu.be" in video_url:
                                            yt_id_match = re.search(r'(?:v=|\/)([0-9A-Za-z_-]{11})', video_url)
                                            if yt_id_match:
                                                video_id = yt_id_match.group(1)
                                                autoplay_html = f"""
                                                    <iframe width="100%" height="315" 
                                                        src="https://www.youtube.com/embed/{video_id}?autoplay=1&mute=0" 
                                                        frameborder="0" 
                                                        allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture" 
                                                        allowfullscreen>
                                                    </iframe>
                                                """
                                                st.components.v1.html(autoplay_html, height=315)
                                                st.success(f" Automatikus lejátszás: **{results[0].get('title', search_query)}**")
                                            else:
                                                st.video(video_url)
                                                st.success(f" Lejátszás: **{results[0].get('title', search_query)}**")
                                        else:
                                            st.warning("Nem találtam biztonságos YouTube linket ehhez a zenéhez.")
                                    else:
                                        st.warning(f"Nem találtam ilyen zenét: {search_query}")
                            except Exception as e:
                                st.error(f"Hiba a zene keresése közben: {e}")

                    # 4. LATENCY ÉS MENTÉS
                    end_time = time.perf_counter()
                    db_repo.log_latency(end_time - start_time)
                    try:
                        db_repo.log_message(active_chat_user, "assistant", full_response, "text", thread_id=current_tid)
                    except Exception as e:
                        st.error(f"Hiba a naplózás során: {e}")
            
            except Exception as main_error:
                st.error(f"Hiba történt a generálás közben: {main_error}")
                st.session_state.generating = False
            
            finally:
                st.session_state.generating = False